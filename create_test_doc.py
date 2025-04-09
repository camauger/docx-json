import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

from docx_json.core.docx_parser import DocxParser
from docx_json.core.html_renderer import HTMLGenerator
from docx_json.core.processor import DocumentProcessor
from docx_json.models import Component


def create_placeholder_image(filename, width=600, height=400, text="Image placeholder"):
    """Crée une image placeholder avec du texte."""
    # Créer le dossier si nécessaire
    os.makedirs(os.path.dirname(filename), exist_ok=True)

    # Créer une image avec un fond gris clair
    image = Image.new("RGB", (width, height), color=(240, 240, 240))
    draw = ImageDraw.Draw(image)

    # Tracer un cadre
    draw.rectangle(
        [(10, 10), (width - 10, height - 10)], outline=(200, 200, 200), width=2
    )

    # Ajouter du texte
    try:
        # Essayer d'utiliser une police système
        font = ImageFont.truetype("arial.ttf", 30)
    except IOError:
        # Utiliser la police par défaut si arial n'est pas disponible
        font = ImageFont.load_default()

    # Calculer la position du texte pour le centrer
    text_width = draw.textlength(text, font=font)
    text_position = ((width - text_width) // 2, height // 2 - 15)
    draw.text(text_position, text, fill=(100, 100, 100), font=font)

    # Sauvegarder l'image
    image.save(filename)


def create_test_document() -> None:
    # Créer des images placeholder pour le carousel
    create_placeholder_image(
        "test_templates/images/placeholder1.jpg", text="Diapositive 1"
    )
    create_placeholder_image(
        "test_templates/images/placeholder2.jpg", text="Diapositive 2"
    )

    # Créer un nouveau document
    doc = Document()

    # Ajouter un titre
    title: Paragraph = doc.add_heading(
        "Document de Test - Styles et Composants Bootstrap", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section pour les styles de texte
    doc.add_heading("Styles de texte", level=2)

    # Titres
    for i in range(1, 7):
        doc.add_heading(f"Titre de niveau {i}", level=i)

    # Styles de paragraphe
    doc.add_heading("Styles de paragraphe", level=2)

    # Paragraphe normal
    p: Paragraph = doc.add_paragraph("Paragraphe normal")

    # Paragraphe en gras
    p = doc.add_paragraph()
    run = p.add_run("Paragraphe en gras")
    run.bold = True

    # Paragraphe en italique
    p = doc.add_paragraph()
    run = p.add_run("Paragraphe en italique")
    run.italic = True

    # Paragraphe souligné
    p = doc.add_paragraph()
    run = p.add_run("Paragraphe souligné")
    run.underline = True

    # Paragraphe avec styles combinés
    p = doc.add_paragraph()
    run = p.add_run("Paragraphe en gras et italique")
    run.bold = True
    run.italic = True

    # Listes
    doc.add_heading("Listes", level=2)

    # Liste à puces simple
    doc.add_heading("Liste à puces simple", level=3)
    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = "List Bullet"
        p.paragraph_format.left_indent = Pt(36)
        p.add_run(f"Élément {i}")

    # Liste numérotée simple
    doc.add_heading("Liste numérotée simple", level=3)
    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = "List Number"
        p.paragraph_format.left_indent = Pt(36)
        p.add_run(f"Élément {i}")

    # Liste à puces avec styles
    doc.add_heading("Liste à puces avec styles", level=3)
    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = "List Bullet"
        p.paragraph_format.left_indent = Pt(36)
        run = p.add_run(f"Élément {i}")
        if i == 1:
            run.bold = True
        elif i == 2:
            run.italic = True
        elif i == 3:
            run.underline = True

    # Liste imbriquée
    doc.add_heading("Liste imbriquée", level=3)
    # Premier niveau
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.paragraph_format.left_indent = Pt(36)
    p.add_run("Élément principal 1")

    # Sous-éléments
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.paragraph_format.left_indent = Pt(72)
    p.add_run("Sous-élément 1.1")

    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.paragraph_format.left_indent = Pt(72)
    p.add_run("Sous-élément 1.2")

    # Deuxième élément principal
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.paragraph_format.left_indent = Pt(36)
    p.add_run("Élément principal 2")

    # Section pour les composants Bootstrap
    doc.add_heading("Composants Bootstrap", level=2)

    # Accordéon
    doc.add_heading("Accordéon", level=3)
    doc.add_paragraph("[Accordéon]")
    doc.add_heading("Axe 1 - Premier axe", level=4)
    doc.add_paragraph("Contenu du premier axe de l'accordéon.")
    doc.add_heading("Axe 2 - Deuxième axe", level=4)
    doc.add_paragraph("Contenu du deuxième axe de l'accordéon.")
    doc.add_heading("Axe 3 - Troisième axe", level=4)
    doc.add_paragraph("Contenu du troisième axe de l'accordéon.")
    doc.add_paragraph("[Fin Accordéon]")

    # Onglets
    doc.add_heading("Onglets", level=3)
    doc.add_paragraph("[Onglets]")
    doc.add_heading("Onglet 1", level=4)
    doc.add_paragraph("Contenu du premier onglet.")
    doc.add_heading("Onglet 2", level=4)
    doc.add_paragraph("Contenu du deuxième onglet.")
    doc.add_paragraph("[Fin Onglets]")

    # Carousel
    doc.add_heading("Carrousel", level=3)
    doc.add_paragraph("[Carrousel]")

    # Diapositive 1 avec image
    doc.add_heading("Diapositive 1", level=4)
    doc.add_paragraph(
        "Au terme de son parcours scolaire, Thomas accède à une formation technique au niveau collégial. Son intérêt marqué pour les mathématiques et les sciences l’amène à choisir une technique de design industriel. Il envisage aussi de poursuivre des études supérieures en génie industriel. Dans ses loisirs, il participe à des jeux de rôle en ligne et fait partie d’un club de lecture traitant de faits historiques et de sciences."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ajouter une image à la première diapositive
    run = p.add_run()
    run.add_picture("test_templates/images/placeholder1.jpg", width=Pt(300))

    # Diapositive 2 avec image
    doc.add_heading("Diapositive 2", level=4)
    doc.add_paragraph(
        "Au terme de son parcours scolaire, Maxim intègre, à raison de deux jours par semaine, un plateau supervisé en entreprise adaptée dans son quartier. Il participe à des fins de semaine en camp adapté, à des camps de jour pendant les vacances estivales et à des activités d’intégration une fois par semaine au sein d’un organisme communautaire."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ajouter une image à la deuxième diapositive
    run = p.add_run()
    run.add_picture("./test_templates/images/placeholder2.jpg", width=Pt(300))

    doc.add_paragraph("[Fin Carrousel]")

    # Audio
    doc.add_heading("Audio", level=3)
    doc.add_paragraph("Titre de l'audio")
    doc.add_paragraph("[Audio]")
    doc.add_paragraph("[Fin Audio]")
    doc.add_paragraph("Description de l'audio")

    # Vidéo
    doc.add_heading("Vidéo Vimeo", level=3)
    # Utiliser un seul paragraphe avec le format vidéo et l'ID intégré

    doc.add_paragraph("Titre de la vidéo Vimeo")
    doc.add_paragraph(
        "Cette vidéo est hébergée sur la plateforme Vimeo et intégrée directement dans le document."
    )
    doc.add_paragraph("[Vidéo video_id='1069341210']")
    doc.add_paragraph("[Fin Vidéo]")

    # Sauvegarder le document
    doc.save("test_document.docx")

    # Générer le HTML
    parser = DocxParser("test_document.docx", ".")
    elements = parser.parse()

    # Traiter les instructions et les composants
    processed_elements = DocumentProcessor.process_instructions(elements)

    # Imprimer le nombre d'éléments pour déboguer
    print(f"Nombre d'éléments avant traitement: {len(elements)}")
    print(f"Nombre d'éléments après traitement: {len(processed_elements)}")

    # Vérifier les types d'éléments après traitement
    component_count = 0
    for element in processed_elements:
        if element.type == "component" and isinstance(element, Component):
            component_count += 1
            print(
                f"Composant trouvé: {element.component_type} avec {len(element.content)} éléments"
            )

    print(f"Nombre de composants trouvés: {component_count}")

    # Créer la structure JSON pour le générateur HTML
    json_data = {
        "meta": {"title": "Document de Test - Bootstrap Components"},
        "content": [element.to_dict() for element in processed_elements],
        "images": parser.get_images(),
    }

    # Définir un CSS personnalisé pour améliorer l'apparence
    custom_css = """
    body {
      background-color: #f8f9fa;
      color: #212529;
    }
    .container {
      max-width: 900px;
      background-color: #fff;
      box-shadow: 0 0.5rem 1rem rgba(0, 0, 0, 0.15);
      border-radius: 0.5rem;
      padding: 2rem;
      margin-top: 2rem;
      margin-bottom: 2rem;
    }
    h1 {
      color: #0d6efd;
      border-bottom: 2px solid #dee2e6;
      padding-bottom: 0.5rem;
      margin-bottom: 1.5rem;
    }
    h2 {
      color: #198754;
      margin-top: 2rem;
      border-left: 4px solid #198754;
      padding-left: 0.5rem;
    }
    h3 {
      color: #6610f2;
      margin-top: 1.5rem;
    }
    .accordion {
      margin-bottom: 2rem;
    }
    .accordion-button:not(.collapsed) {
      background-color: #e7f1ff;
      color: #0c63e4;
    }
    .nav-tabs .nav-link.active {
      font-weight: bold;
    }
    .carousel {
      margin-bottom: 2rem;
      border: 1px solid #dee2e6;
      border-radius: 0.375rem;
      overflow: hidden;
    }
    .carousel-content {
      min-height: 200px;
    }
    ul, ol {
      background-color: #f8f9fa;
      padding: 1rem 1rem 1rem 2.5rem;
      border-radius: 0.375rem;
    }
    .ratio-16x9 {
      border-radius: 0.375rem;
      overflow: hidden;
      box-shadow: 0 0.25rem 0.5rem rgba(0, 0, 0, 0.1);
    }
    audio {
      box-shadow: 0 0.25rem 0.5rem rgba(0, 0, 0, 0.1);
      border-radius: 0.375rem;
    }
    """

    generator = HTMLGenerator(json_data)
    html_content = generator.generate(custom_css=custom_css)

    # Sauvegarder le HTML
    with open("test_document.html", "w", encoding="utf-8") as f:
        f.write(html_content)


if __name__ == "__main__":
    create_test_document()
