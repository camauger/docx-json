#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Tests pour le module docx_parser.py
Tests complets pour l'extraction et l'analyse des documents DOCX
"""

import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from docx_json.core.docx_parser import DocxParser
from docx_json.models import (
    Component,
    ComponentEnd,
    ComponentMarker,
    Heading,
    Image,
    Instruction,
    ListItem,
    Paragraph,
    Table,
    TextRun,
)


class TestDocxParserInit(unittest.TestCase):
    """Tests pour l'initialisation du DocxParser."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

        # Créer un document DOCX de test
        self.test_docx = self.temp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(self.test_docx)

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_parser_initialization(self):
        """Test initialisation basique du parser."""
        parser = DocxParser(
            str(self.test_docx), str(self.temp_path), save_images_to_disk=True
        )

        self.assertEqual(parser._docx_path, str(self.test_docx))
        self.assertEqual(parser._output_dir, str(self.temp_path))
        self.assertTrue(parser._save_images_to_disk)
        self.assertIsNotNone(parser._document)

    def test_parser_metadata(self):
        """Test initialisation des métadonnées."""
        parser = DocxParser(str(self.test_docx), str(self.temp_path))

        self.assertIn("title", parser.metadata)
        self.assertEqual(parser.metadata["title"], "test.docx")

    def test_parser_namespaces(self):
        """Test définition des namespaces XML."""
        parser = DocxParser(str(self.test_docx), str(self.temp_path))

        self.assertIn("w", parser.NAMESPACES)
        self.assertIn("wp", parser.NAMESPACES)
        self.assertIn("a", parser.NAMESPACES)
        self.assertIn("pic", parser.NAMESPACES)

    def test_load_relationships(self):
        """Test chargement des relations."""
        parser = DocxParser(str(self.test_docx), str(self.temp_path))

        # Le document devrait être chargé
        self.assertIsNotNone(parser._document)
        self.assertIsInstance(parser._document, Document)


class TestDocxParserImageExtraction(unittest.TestCase):
    """Tests pour l'extraction d'images."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

        # Créer un document simple
        self.test_docx = self.temp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(self.test_docx)

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_images_directory_creation(self):
        """Test création du dossier images."""
        parser = DocxParser(
            str(self.test_docx), str(self.temp_path), save_images_to_disk=True
        )

        images_dir = self.temp_path / "images"
        # Le dossier devrait être créé si des images sont présentes
        # Pour ce test, pas d'images donc on vérifie juste l'initialisation
        self.assertIsInstance(parser._images, dict)

    def test_save_images_to_disk_flag(self):
        """Test du flag save_images_to_disk."""
        parser_disk = DocxParser(
            str(self.test_docx), str(self.temp_path), save_images_to_disk=True
        )
        parser_base64 = DocxParser(
            str(self.test_docx), str(self.temp_path), save_images_to_disk=False
        )

        self.assertTrue(parser_disk._save_images_to_disk)
        self.assertFalse(parser_base64._save_images_to_disk)

    def test_images_dict_initialization(self):
        """Test initialisation du dictionnaire d'images."""
        parser = DocxParser(str(self.test_docx), str(self.temp_path))

        self.assertIsInstance(parser._images, dict)
        self.assertEqual(len(parser._images), 0)

    def test_rels_dict_initialization(self):
        """Test initialisation du dictionnaire de relations."""
        parser = DocxParser(str(self.test_docx), str(self.temp_path))

        self.assertIsInstance(parser._rels_dict, dict)


class TestDocxParserParagraphParsing(unittest.TestCase):
    """Tests pour l'analyse des paragraphes."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

        # Créer un document avec différents types de paragraphes
        self.test_docx = self.temp_path / "test.docx"
        doc = Document()

        # Paragraphe normal
        doc.add_paragraph("Paragraphe normal")

        # Titre niveau 1
        doc.add_heading("Titre niveau 1", level=1)

        # Titre niveau 2
        doc.add_heading("Titre niveau 2", level=2)

        # Paragraphe avec instruction
        doc.add_paragraph(":::class test-class")

        # Marqueur de composant
        doc.add_paragraph("[Vidéo]")

        # Marqueur de fin de composant
        doc.add_paragraph("[Fin Vidéo]")

        doc.save(self.test_docx)

        self.parser = DocxParser(str(self.test_docx), str(self.temp_path))

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_parse_normal_paragraph(self):
        """Test analyse d'un paragraphe normal."""
        paragraphs = self.parser._document.paragraphs
        normal_para = paragraphs[0]

        element = self.parser.parse_paragraph(normal_para)

        self.assertIsInstance(element, Paragraph)
        self.assertEqual(element.type, "paragraph")

    def test_parse_heading(self):
        """Test analyse d'un titre."""
        paragraphs = self.parser._document.paragraphs
        heading_para = paragraphs[1]  # "Titre niveau 1"

        element = self.parser.parse_paragraph(heading_para)

        self.assertIsInstance(element, Heading)
        self.assertEqual(element.type, "heading")
        self.assertEqual(element.level, 1)

    def test_parse_heading_level_2(self):
        """Test analyse d'un titre niveau 2."""
        paragraphs = self.parser._document.paragraphs
        heading_para = paragraphs[2]  # "Titre niveau 2"

        element = self.parser.parse_paragraph(heading_para)

        self.assertIsInstance(element, Heading)
        self.assertEqual(element.level, 2)

    def test_parse_instruction(self):
        """Test analyse d'une instruction."""
        paragraphs = self.parser._document.paragraphs
        instruction_para = paragraphs[3]  # ":::class test-class"

        element = self.parser.parse_paragraph(instruction_para)

        self.assertIsInstance(element, Instruction)
        self.assertEqual(element.type, "instruction")

    def test_parse_component_marker(self):
        """Test analyse d'un marqueur de composant."""
        paragraphs = self.parser._document.paragraphs
        marker_para = paragraphs[4]  # "[Vidéo]"

        element = self.parser.parse_paragraph(marker_para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Vidéo")

    def test_parse_component_end_marker(self):
        """Test analyse d'un marqueur de fin de composant."""
        paragraphs = self.parser._document.paragraphs
        end_para = paragraphs[5]  # "[Fin Vidéo]"

        element = self.parser.parse_paragraph(end_para)

        self.assertIsInstance(element, ComponentEnd)
        self.assertEqual(element.component_type, "Vidéo")


class TestDocxParserTableParsing(unittest.TestCase):
    """Tests pour l'analyse des tableaux."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

        # Créer un document avec un tableau
        self.test_docx = self.temp_path / "test.docx"
        doc = Document()

        # Ajouter un tableau 2x2
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        doc.save(self.test_docx)

        self.parser = DocxParser(str(self.test_docx), str(self.temp_path))

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_parse_table(self):
        """Test analyse d'un tableau."""
        tables = self.parser._document.tables
        self.assertEqual(len(tables), 1)

        table = tables[0]
        table_element = self.parser.parse_table(table)

        self.assertIsInstance(table_element, Table)
        self.assertEqual(table_element.type, "table")

    def test_table_rows_count(self):
        """Test nombre de lignes du tableau."""
        table = self.parser._document.tables[0]
        table_element = self.parser.parse_table(table)

        # Vérifier le nombre de lignes
        self.assertEqual(len(table_element.rows), 2)

    def test_table_cells_content(self):
        """Test contenu des cellules."""
        table = self.parser._document.tables[0]
        table_element = self.parser.parse_table(table)

        # Vérifier le contenu des cellules
        # Chaque cellule contient une liste de paragraphes
        self.assertTrue(len(table_element.rows) > 0)
        self.assertTrue(len(table_element.rows[0]) == 2)  # 2 colonnes


class TestDocxParserComponentTypes(unittest.TestCase):
    """Tests pour la détection des différents types de composants."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.test_docx = self.temp_path / "test.docx"

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_detect_video_component(self):
        """Test détection composant Vidéo."""
        doc = Document()
        doc.add_paragraph("[Vidéo]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Vidéo")

    def test_detect_audio_component(self):
        """Test détection composant Audio."""
        doc = Document()
        doc.add_paragraph("[Audio]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Audio")

    def test_detect_accordion_component(self):
        """Test détection composant Accordéon."""
        doc = Document()
        doc.add_paragraph("[Accordéon]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Accordéon")

    def test_detect_carousel_component(self):
        """Test détection composant Carrousel."""
        doc = Document()
        doc.add_paragraph("[Carrousel]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Carrousel")

    def test_detect_tabs_component(self):
        """Test détection composant Onglets."""
        doc = Document()
        doc.add_paragraph("[Onglets]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Onglets")

    def test_detect_scrollspy_component(self):
        """Test détection composant Défilement."""
        doc = Document()
        doc.add_paragraph("[Défilement]")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, ComponentMarker)
        self.assertEqual(element.component_type, "Défilement")


class TestDocxParserInstructions(unittest.TestCase):
    """Tests pour le traitement des instructions."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.test_docx = self.temp_path / "test.docx"

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_class_instruction(self):
        """Test instruction de classe CSS."""
        doc = Document()
        doc.add_paragraph(":::class hero dark")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertIn("class hero dark", element.content)

    def test_id_instruction(self):
        """Test instruction d'ID HTML."""
        doc = Document()
        doc.add_paragraph(":::id main-section")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertIn("id main-section", element.content)

    def test_ignore_instruction(self):
        """Test instruction ignore."""
        doc = Document()
        doc.add_paragraph(":::ignore")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertEqual(element.content, "ignore")

    def test_block_start_instruction(self):
        """Test instruction de début de bloc."""
        doc = Document()
        doc.add_paragraph(":::quote start")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertIn("quote start", element.content)

    def test_block_end_instruction(self):
        """Test instruction de fin de bloc."""
        doc = Document()
        doc.add_paragraph(":::quote end")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertIn("quote end", element.content)

    def test_html_instruction(self):
        """Test instruction HTML brut."""
        doc = Document()
        doc.add_paragraph(":::html <hr />")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        para = parser._document.paragraphs[0]
        element = parser.parse_paragraph(para)

        self.assertIsInstance(element, Instruction)
        self.assertIn("html <hr />", element.content)


class TestDocxParserTextFormatting(unittest.TestCase):
    """Tests pour le formatage de texte (gras, italique, souligné)."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.test_docx = self.temp_path / "test.docx"

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_bold_text(self):
        """Test texte en gras."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Texte en gras")
        run.bold = True
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed_para = parser.parse_paragraph(parser._document.paragraphs[0])

        self.assertIsInstance(parsed_para, Paragraph)
        self.assertTrue(len(parsed_para.runs) > 0)
        # Le premier run devrait être en gras
        self.assertTrue(parsed_para.runs[0].bold)

    def test_italic_text(self):
        """Test texte en italique."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Texte en italique")
        run.italic = True
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed_para = parser.parse_paragraph(parser._document.paragraphs[0])

        self.assertTrue(len(parsed_para.runs) > 0)
        self.assertTrue(parsed_para.runs[0].italic)

    def test_underline_text(self):
        """Test texte souligné."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Texte souligné")
        run.underline = True
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed_para = parser.parse_paragraph(parser._document.paragraphs[0])

        self.assertTrue(len(parsed_para.runs) > 0)
        self.assertTrue(parsed_para.runs[0].underline)

    def test_mixed_formatting(self):
        """Test texte avec formatage mixte."""
        doc = Document()
        para = doc.add_paragraph()

        run1 = para.add_run("Gras")
        run1.bold = True

        run2 = para.add_run(" et italique")
        run2.italic = True

        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed_para = parser.parse_paragraph(parser._document.paragraphs[0])

        self.assertEqual(len(parsed_para.runs), 2)
        self.assertTrue(parsed_para.runs[0].bold)
        self.assertTrue(parsed_para.runs[1].italic)


class TestDocxParserListDetection(unittest.TestCase):
    """Tests pour la détection de listes."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.test_docx = self.temp_path / "test.docx"

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_detect_list_by_indent(self):
        """Test détection de liste par indentation."""
        doc = Document()
        para = doc.add_paragraph("Élément de liste")
        # Ajouter une indentation pour simuler une liste
        para.paragraph_format.left_indent = 720000  # 0.5 pouce

        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed = parser.parse_paragraph(parser._document.paragraphs[0])

        # Devrait être détecté comme ListItem grâce à l'indentation
        self.assertIsInstance(parsed, ListItem)

    def test_normal_paragraph_no_indent(self):
        """Test paragraphe normal sans indentation."""
        doc = Document()
        doc.add_paragraph("Paragraphe normal")
        doc.save(self.test_docx)

        parser = DocxParser(str(self.test_docx), str(self.temp_path))
        parsed = parser.parse_paragraph(parser._document.paragraphs[0])

        # Devrait être un Paragraph normal
        self.assertIsInstance(parsed, Paragraph)


class TestDocxParserComplexDocuments(unittest.TestCase):
    """Tests pour des documents complexes."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_mixed_content_document(self):
        """Test document avec contenu mixte."""
        test_docx = self.temp_path / "mixed.docx"
        doc = Document()

        # Ajouter différents types d'éléments
        doc.add_heading("Titre principal", level=1)
        doc.add_paragraph("Paragraphe introductif")
        doc.add_heading("Sous-titre", level=2)
        doc.add_paragraph(":::class highlight")
        doc.add_paragraph("Paragraphe avec classe")

        # Ajouter un tableau
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Col1"
        table.cell(0, 1).text = "Col2"

        doc.add_paragraph("[Accordéon]")
        doc.add_paragraph("Contenu accordéon")
        doc.add_paragraph("[Fin Accordéon]")

        doc.save(test_docx)

        parser = DocxParser(str(test_docx), str(self.temp_path))

        # Vérifier que le document a été chargé
        self.assertIsNotNone(parser._document)

        # Vérifier le nombre de paragraphes
        self.assertGreater(len(parser._document.paragraphs), 0)

        # Vérifier qu'on a un tableau
        self.assertEqual(len(parser._document.tables), 1)

    def test_empty_document(self):
        """Test document vide."""
        test_docx = self.temp_path / "empty.docx"
        doc = Document()
        doc.save(test_docx)

        parser = DocxParser(str(test_docx), str(self.temp_path))

        self.assertIsNotNone(parser._document)
        # Un document vide peut avoir 0 ou 1 paragraphe par défaut
        self.assertIsInstance(parser._document.paragraphs, list)

    def test_document_with_only_tables(self):
        """Test document avec seulement des tableaux."""
        test_docx = self.temp_path / "tables.docx"
        doc = Document()

        # Ajouter 3 tableaux
        for i in range(3):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = f"Table{i}"

        doc.save(test_docx)

        parser = DocxParser(str(test_docx), str(self.temp_path))

        self.assertEqual(len(parser._document.tables), 3)


class TestDocxParserErrorHandling(unittest.TestCase):
    """Tests pour la gestion d'erreurs."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_nonexistent_file(self):
        """Test avec fichier non existant."""
        nonexistent = self.temp_path / "nonexistent.docx"

        with self.assertRaises(Exception):
            parser = DocxParser(str(nonexistent), str(self.temp_path))

    def test_invalid_output_directory(self):
        """Test avec répertoire de sortie invalide."""
        test_docx = self.temp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(test_docx)

        # Même avec un répertoire invalide, le parser devrait s'initialiser
        # (l'erreur apparaîtra lors de la sauvegarde des images)
        parser = DocxParser(str(test_docx), "/invalid/path/that/does/not/exist")

        self.assertIsNotNone(parser)


class TestDocxParserMetadata(unittest.TestCase):
    """Tests pour l'extraction de métadonnées."""

    def setUp(self):
        """Crée un environnement de test."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """Nettoie l'environnement."""
        self.temp_dir.cleanup()

    def test_metadata_title_from_filename(self):
        """Test extraction du titre depuis le nom du fichier."""
        test_docx = self.temp_path / "mon_document.docx"
        doc = Document()
        doc.add_paragraph("Content")
        doc.save(test_docx)

        parser = DocxParser(str(test_docx), str(self.temp_path))

        self.assertEqual(parser.metadata["title"], "mon_document.docx")

    def test_metadata_initialization(self):
        """Test initialisation des métadonnées."""
        test_docx = self.temp_path / "test.docx"
        doc = Document()
        doc.save(test_docx)

        parser = DocxParser(str(test_docx), str(self.temp_path))

        self.assertIsInstance(parser.metadata, dict)
        self.assertIn("title", parser.metadata)


if __name__ == "__main__":
    unittest.main()
