#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Tests unitaires pour la classe DocxConverter et les modèles de données
"""

import json
import os
import tempfile
import unittest

from docx import Document

from docx_json.core.converter import DocxConverter, HTMLGenerator
from docx_json.models.elements import (
    Component,
    DocumentElement,
    Heading,
    Image,
    ListItem,
    Paragraph,
    Table,
    TextRun,
)


class TestElementModels(unittest.TestCase):
    """Tests unitaires pour les modèles de données."""

    def test_text_run(self):
        """Teste la création et la sérialisation d'un TextRun."""
        run = TextRun(text="Test", bold=True, italic=False, underline=True)

        self.assertEqual(run.text, "Test")
        self.assertTrue(run.bold)
        self.assertFalse(run.italic)
        self.assertTrue(run.underline)

        # Test de la sérialisation
        run_dict = run.to_dict()
        self.assertEqual(run_dict["text"], "Test")
        self.assertTrue(run_dict["bold"])
        self.assertFalse(run_dict["italic"])
        self.assertTrue(run_dict["underline"])

    def test_paragraph(self):
        """Teste la création et la sérialisation d'un Paragraph."""
        para = Paragraph()
        para.add_run(TextRun(text="Premier segment", bold=True))
        para.add_run(TextRun(text=" et deuxième segment", italic=True))

        # Vérifier que les runs ont bien été ajoutés
        self.assertEqual(len(para.runs), 2)

        # Test de la sérialisation
        para_dict = para.to_dict()
        self.assertEqual(para_dict["type"], "paragraph")
        self.assertEqual(len(para_dict["runs"]), 2)
        self.assertEqual(para_dict["runs"][0]["text"], "Premier segment")
        self.assertTrue(para_dict["runs"][0]["bold"])
        self.assertEqual(para_dict["runs"][1]["text"], " et deuxième segment")
        self.assertTrue(para_dict["runs"][1]["italic"])

    def test_heading(self):
        """Teste la création et la sérialisation d'un Heading."""
        heading = Heading(level=2)
        heading.add_run(TextRun(text="Titre de niveau 2", bold=True))

        # Test de la sérialisation
        heading_dict = heading.to_dict()
        self.assertEqual(heading_dict["type"], "heading")
        self.assertEqual(heading_dict["level"], 2)
        self.assertEqual(len(heading_dict["runs"]), 1)
        self.assertEqual(heading_dict["runs"][0]["text"], "Titre de niveau 2")

    def test_component(self):
        """Teste la création et la sérialisation d'un Component."""
        component = Component(component_type="Accordéon")

        # Ajouter des éléments au composant
        heading = Heading(level=3)
        heading.add_run(TextRun(text="Section d'accordéon", bold=True))

        para = Paragraph()
        para.add_run(TextRun(text="Contenu de l'accordéon"))

        component.add_element(heading)
        component.add_element(para)

        # Test de la sérialisation
        comp_dict = component.to_dict()
        self.assertEqual(comp_dict["type"], "component")
        self.assertEqual(comp_dict["component_type"], "Accordéon")
        self.assertEqual(len(comp_dict["content"]), 2)
        self.assertEqual(comp_dict["content"][0]["type"], "heading")
        self.assertEqual(comp_dict["content"][1]["type"], "paragraph")


class TestDocxConverter(unittest.TestCase):
    """Tests unitaires pour la classe DocxConverter."""

    def setUp(self):
        """Initialise l'environnement de test."""
        # Créer un document temporaire pour les tests
        self.test_doc = Document()
        self.test_doc.add_heading("Titre de test", level=1)
        self.test_doc.add_paragraph("Ceci est un paragraphe de test.")

        # Ajouter un paragraphe avec une instruction
        self.test_doc.add_paragraph(":::class test-class")
        self.test_doc.add_paragraph("Paragraphe avec classe")

        # Créer un tableau simple
        table = self.test_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        # Enregistrer le document
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = os.path.join(self.temp_dir.name, "test.docx")
        self.test_doc.save(self.temp_path)

        # Initialiser le convertisseur
        self.converter = DocxConverter(self.temp_path, self.temp_dir.name)
        # S'assurer que le document est chargé
        self.converter.load_document()

    def tearDown(self):
        """Nettoie l'environnement après les tests."""
        self.temp_dir.cleanup()

    def test_load_document(self):
        """Teste le chargement d'un document DOCX."""
        self.assertIsNotNone(self.converter._document)

    def test_parse_paragraph(self):
        """Teste l'analyse d'un paragraphe."""
        # Vérifier que le document est chargé
        if self.converter._document is None:
            self.skipTest("Le document n'est pas chargé")
            return

        # Tester l'analyse d'un titre
        if (
            not hasattr(self.converter._document, "paragraphs")
            or not self.converter._document.paragraphs
        ):
            self.skipTest("Le document n'a pas de paragraphes")
            return

        heading = self.converter._document.paragraphs[0]
        element = self.converter.parse_paragraph(heading)

        # Vérifier que c'est bien un Heading
        self.assertIsInstance(element, Heading)
        # Accéder aux attributs spécifiques à Heading
        if isinstance(element, Heading):
            self.assertEqual(element.level, 1)
            self.assertEqual(len(element.runs), 1)
            self.assertEqual(element.runs[0].text, "Titre de test")

        # Tester l'analyse d'un paragraphe normal
        para = self.converter._document.paragraphs[1]
        element = self.converter.parse_paragraph(para)

        self.assertIsInstance(element, Paragraph)
        # Accéder aux attributs spécifiques à Paragraph
        if isinstance(element, Paragraph):
            self.assertEqual(len(element.runs), 1)
            self.assertEqual(element.runs[0].text, "Ceci est un paragraphe de test.")

    def test_process_instructions(self):
        """Teste le traitement des instructions intégrées."""
        # Vérifier que le document est chargé
        if self.converter._document is None:
            self.skipTest("Le document n'est pas chargé")
            return

        # Vérifier qu'il y a des paragraphes
        if (
            not hasattr(self.converter._document, "paragraphs")
            or not self.converter._document.paragraphs
        ):
            self.skipTest("Le document n'a pas de paragraphes")
            return

        # Extraire et convertir tous les éléments
        elements = []
        for para in self.converter._document.paragraphs:
            elements.append(self.converter.parse_paragraph(para))

        # Traiter les instructions
        processed = self.converter.process_instructions(elements)

        # Vérifier que l'instruction a été correctement appliquée
        # 3 au lieu de 4 car l'instruction est retirée
        self.assertEqual(len(processed), 3)
        self.assertEqual(processed[2].html_class, "test-class")

    def test_convert(self):
        """Teste la conversion complète du document."""
        json_data = self.converter.convert()

        # Vérifier la structure du JSON
        self.assertIn("meta", json_data)
        self.assertIn("content", json_data)
        self.assertIn("images", json_data)

        # Vérifier que le contenu a été correctement analysé
        self.assertEqual(len(json_data["content"]), 3)
        self.assertEqual(json_data["content"][0]["type"], "heading")
        self.assertEqual(json_data["content"][1]["type"], "paragraph")
        self.assertEqual(json_data["content"][2]["type"], "paragraph")
        self.assertEqual(json_data["content"][2]["html_class"], "test-class")


class TestHTMLGenerator(unittest.TestCase):
    """Tests unitaires pour la classe HTMLGenerator."""

    def setUp(self):
        """Initialise les données de test."""
        self.json_data = {
            "meta": {"title": "test.docx"},
            "content": [
                {
                    "type": "heading",
                    "level": 1,
                    "runs": [{"text": "Titre", "bold": True}],
                },
                {"type": "paragraph", "runs": [{"text": "Paragraphe", "bold": False}]},
                {
                    "type": "component",
                    "component_type": "Accordéon",
                    "content": [
                        {
                            "type": "heading",
                            "level": 2,
                            "runs": [{"text": "Section Accordéon", "bold": True}],
                        },
                        {
                            "type": "paragraph",
                            "runs": [{"text": "Contenu caché", "bold": False}],
                        },
                    ],
                },
            ],
            "images": {},
        }

        self.generator = HTMLGenerator(self.json_data)

    def test_generate_html_structure(self):
        """Teste la génération de la structure HTML de base."""
        # Appeler la méthode correcte: generate() au lieu de generate_html()
        html = self.generator.generate()

        # Vérifier la structure HTML
        self.assertTrue("<!DOCTYPE html>" in html)
        self.assertTrue('<html lang="fr">' in html)
        self.assertTrue("<head>" in html)
        self.assertTrue("<body>" in html)
        # Vérifie que Bootstrap est chargé
        self.assertTrue("bootstrap" in html)

    def test_generate_element_html(self):
        """Teste la génération HTML pour différents types d'éléments."""
        # Appeler la méthode correcte: _generate_element_html au lieu de generate_element_html

        # Tester un titre
        heading_html = "\n".join(
            self.generator._generate_element_html(self.json_data["content"][0])
        )
        self.assertTrue("<h1>" in heading_html)
        self.assertTrue("<strong>Titre</strong>" in heading_html)

        # Tester un paragraphe
        para_html = "\n".join(
            self.generator._generate_element_html(self.json_data["content"][1])
        )
        self.assertTrue("<p>" in para_html)
        self.assertTrue("Paragraphe" in para_html)

        # Tester un composant accordéon
        accordion_html = "\n".join(
            self.generator._generate_element_html(self.json_data["content"][2])
        )
        self.assertTrue("accordion" in accordion_html.lower())
        self.assertTrue("accordion-button" in accordion_html)
        self.assertTrue("accordion-body" in accordion_html)
        self.assertTrue("Section Accordéon" in accordion_html)
        self.assertTrue("Contenu caché" in accordion_html)


if __name__ == "__main__":
    unittest.main()
