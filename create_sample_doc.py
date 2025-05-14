#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Script pour créer un document Word formaté pour être converti
par docx-json en JSON et HTML.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_sample_document():
    # Créer un nouveau document Word
    doc = Document()

    # --- TITRE ET INSTRUCTIONS SPÉCIALES ---
    # Titre avec instruction spéciale de classe
    p = doc.add_paragraph(":::class hero")
    p = doc.add_paragraph("Titre principal")

    # Exemple de commentaire qui sera filtré
    p = doc.add_paragraph(
        "### Ce commentaire ne sera pas visible dans le résultat final ###"
    )

    # Contenu avec ID spécifique
    p = doc.add_paragraph(":::id unique-section")
    p = doc.add_paragraph("Contenu spécifique")

    # Citation avec instruction blockquote
    p = doc.add_paragraph(":::quote start")
    p = doc.add_paragraph("Citation importante")
    p = doc.add_paragraph(":::quote end")

    # --- VIDÉO ---
    h = doc.add_heading("Titre de la vidéo", level=2)

    p = doc.add_paragraph("[Vidéo]")
    p = doc.add_paragraph(
        "Pour mieux comprendre la différenciation pédagogique, nous vous invitons à visionner la capsule suivante, préparée par la Direction de l'adaptation scolaire."
    )
    p = doc.add_paragraph()
    p.add_run("Cliquez sur le centre de l'image pour visionner la capsule.")
    p = doc.add_paragraph("[Fin Vidéo]")

    p = doc.add_paragraph("Transcription ")
    p = doc.add_paragraph("[Accordéon]")
    p = doc.add_heading("La différenciation pédagogique", level=3)
    p = doc.add_paragraph(
        "Il n'y a pas de doute que les besoins des élèves sont à la fois semblables et différents, qu'ils évoluent au fil du temps et qu'ils sont susceptibles d'influencer l'apprentissage. La nécessité de différencier l'enseignement prend son origine dans le fait qu'il y a une très grande hétérogénéité d'élèves qui compose les classes."
    )
    p = doc.add_paragraph("[Fin Accordéon]")

    # --- AUDIO ---
    h = doc.add_heading("Titre de l'audio", level=2)

    p = doc.add_paragraph("[Audio]")
    p = doc.add_paragraph(
        "Découvrez les explications de Marie en écoutant l'extrait suivant."
    )
    p = doc.add_paragraph()
    p.add_run("Marie, représentante du SASIE")
    p = doc.add_paragraph("[Fin Audio]")

    p = doc.add_paragraph("Transcription ")
    p = doc.add_paragraph("[Accordéon]")
    p = doc.add_heading("Une rupture complète", level=3)
    p = doc.add_paragraph(
        "Au Québec, l'adoption internationale est dite plénière, c'est-à-dire qu'il y a rupture complète des liens de filiation de l'enfant avec ses parents d'origine et, par conséquent, avec toute sa famille élargie. Une nouvelle filiation est créée avec les parents adoptants."
    )
    p = doc.add_paragraph("[Fin Accordéon]")

    # --- ACCORDÉON AVEC ICÔNE ---
    p = doc.add_paragraph("[Accordéon avec icône]")
    p = doc.add_heading(
        "La complexité d'une tâche a-t-elle une incidence sur l'activité du cerveau?",
        level=3,
    )
    p = doc.add_paragraph(
        "De manière générale, plus une tâche devient difficile ou complexe, plus la coordination de l'activité entre les régions du cerveau impliquées augmente. Par exemple, pour une tâche simple, telle qu'une suite à compléter (1, 2, 3, ?), on observe une activité relativement coordonnée entre les régions frontales (fonctions exécutives) et certaines autres régions."
    )
    p = doc.add_paragraph("[Fin Accordéon]")

    # --- CARROUSEL ---
    p = doc.add_paragraph("[Carrousel]")
    p = doc.add_paragraph(
        "Pour lire la description de Thomas, cliquez sur les flèches de droite et de gauche plus bas."
    )

    p = doc.add_heading("SA FAMILLE", level=3)
    p = doc.add_paragraph("Vit en appartement avec ses parents")
    p = doc.add_paragraph("A une sœur plus âgée")

    p = doc.add_heading("CE QU'IL AIME", level=3)
    p = doc.add_paragraph("Les jeux vidéo")
    p = doc.add_paragraph("Le karaté")
    p = doc.add_paragraph("Les mathématiques")
    p = doc.add_paragraph("[Fin Carrousel]")

    # --- ONGLETS ---
    p = doc.add_paragraph("[Onglets]")
    p = doc.add_paragraph(
        "Pour lire la description de chacun des vecteurs, cliquez sur le vecteur correspondant ou utilisez les flèches de droite et de gauche plus bas."
    )
    p = doc.add_paragraph()
    p.add_run("La mission de l'école québécoise est composée de trois grands vecteurs")

    p = doc.add_heading("Instruire", level=3)
    p = doc.add_paragraph(
        "L'école a comme première responsabilité la formation de l'esprit de chaque élève. Même si elle ne constitue pas le seul lieu d'apprentissage de l'élève, elle joue un rôle irremplaçable en ce qui a trait au développement intellectuel et à l'acquisition de connaissances."
    )

    p = doc.add_heading("Socialiser", level=3)
    p = doc.add_paragraph(
        "Dans une société pluraliste comme la société québécoise, l'école joue un rôle d'agent de cohésion en contribuant à l'apprentissage du vivre-ensemble et au développement d'un sentiment d'appartenance à la collectivité."
    )

    p = doc.add_heading("Qualifier", level=3)
    p = doc.add_paragraph(
        "L'école a le devoir de rendre possible la réussite scolaire de tous les élèves et de faciliter leur intégration sociale et professionnelle, quelle que soit la voie qu'ils choisiront au terme de leur formation."
    )
    p = doc.add_paragraph("[Fin Onglets]")

    # --- DÉFILEMENT AVEC ÉTAPES ---
    p = doc.add_paragraph("[Défilement du menu de gauche avec étapes]")
    p = doc.add_paragraph(
        "Les étapes de réalisation d'un projet d'adoption internationale"
    )

    p = doc.add_heading("Préparation à l'adoption internationale", level=3)
    p = doc.add_paragraph(
        "Le programme d'information et de sensibilisation en ligne L'adoption internationale : les premiers pas de ma réflexion, auquel vous participez actuellement, est la toute première démarche à entreprendre. Il est conçu pour vous aider à prendre une décision éclairée, c'est-à-dire à déterminer si l'adoption internationale est pour vous."
    )

    p = doc.add_heading("Élaboration du projet", level=3)
    p = doc.add_paragraph("Cette deuxième étape consiste à :")
    p = doc.add_paragraph("Identifier les pays d'origine qui vous intéressent.")
    p = doc.add_paragraph(
        "Vérifier les conditions et critères du Québec et des pays d'origine."
    )

    p = doc.add_heading("Ouverture du dossier au SAI", level=3)
    p = doc.add_paragraph(
        "Lorsque vous signerez votre contrat avec l'organisme agréé (OA), celui-ci vous remettra un formulaire de demande d'ouverture d'un dossier d'adoption destiné au Secrétariat aux services internationaux à l'enfant (SASIE)."
    )

    p = doc.add_heading("Évaluation psychosociale", level=3)
    p = doc.add_paragraph(
        "« L'évaluation psychosociale est souvent initialement perçue comme intrusive. Mais après l'avoir complétée, plusieurs candidats à l'adoption en voient les bénéfices. Ils réalisent que cette étape leur permet de mieux réfléchir à leur projet d'adoption et de s'assurer que leurs attentes et leurs besoins seront entendus. »"
    )
    p = doc.add_paragraph("[Fin Défilement]")

    # Sauvegarder le document
    document_path = "synthese_moodle.docx"
    doc.save(document_path)
    print(f"Document créé avec succès: {os.path.abspath(document_path)}")


if __name__ == "__main__":
    create_sample_document()
