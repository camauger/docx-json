#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Crée un document Word avec le nouveau composant Consignes
pour tester son fonctionnement dans la conversion.
"""

import os

from docx import Document


def create_consignes_test_document():
    # Créer un nouveau document Word
    doc = Document()

    # Ajouter l'instruction de titre
    p = doc.add_paragraph(":::title Document de test - Composant Consignes")

    # Ajouter un titre principal
    h = doc.add_heading("Test du composant Consignes", level=1)

    # Ajouter un paragraphe d'introduction
    p = doc.add_paragraph(
        "Ce document teste le nouveau composant 'Consignes' qui permet d'afficher des instructions avec un style spécifique."
    )

    # Ajouter un composant Consignes
    p = doc.add_paragraph("[Consignes]")
    p = doc.add_paragraph(
        "Veuillez lire attentivement les instructions suivantes avant de commencer l'activité:"
    )
    p = doc.add_paragraph("1. Prévoyez environ 30 minutes pour compléter l'exercice.")
    p = doc.add_paragraph("2. Utilisez uniquement les ressources autorisées.")
    p = doc.add_paragraph("3. Soumettez votre travail avant la date limite indiquée.")
    p = doc.add_paragraph("[Fin Consignes]")

    # Ajouter un autre paragraphe normal
    p = doc.add_paragraph(
        "Le texte ci-dessus devrait apparaître avec un fond jaune pâle et une bordure sur la gauche."
    )

    # Ajouter un autre composant Consignes
    p = doc.add_paragraph("[Consignes]")
    h = doc.add_heading("Instructions pour l'exercice 2", level=3)
    p = doc.add_paragraph("Pour cet exercice, vous devrez:")
    p = doc.add_paragraph("- Analyser le document fourni")
    p = doc.add_paragraph("- Répondre aux questions")
    p = doc.add_paragraph("- Justifier vos réponses")
    p = doc.add_paragraph("[Fin Consignes]")

    # Sauvegarder le document
    output_path = "consignes_test.docx"
    doc.save(output_path)
    print(f"Document créé avec succès : {os.path.abspath(output_path)}")


if __name__ == "__main__":
    create_consignes_test_document()
