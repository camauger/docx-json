#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Classe principale de conversion DOCX vers JSON/HTML
"""

import base64
import datetime
import json
import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple, Union, cast

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from docx_json.models.elements import (
    Block,
    Component,
    ComponentEnd,
    ComponentMarker,
    DocumentElement,
    Heading,
    Image,
    Instruction,
    ListItem,
    Paragraph,
    RawHTML,
    Table,
    TextRun,
)


class DocxConverter:
    """
    Classe pour convertir des documents DOCX en JSON et HTML.
    """

    def __init__(
        self, docx_path: str, output_dir: str, save_images_to_disk: bool = True
    ):
        """
        Initialise le convertisseur.

        Args:
            docx_path: Chemin vers le fichier .docx
            output_dir: Répertoire de sortie pour les fichiers générés
            save_images_to_disk: Si True, sauvegarde les images sur disque. Sinon, encode en base64.
        """
        self._docx_path = docx_path
        self._output_dir = output_dir
        self._save_images_to_disk = save_images_to_disk
        self._document = None
        self._images = {}
        self._rels_dict = {}

    def load_document(self) -> None:
        """Charge le document DOCX."""
        logging.info(f"Chargement du document: {self._docx_path}")
        self._document = Document(self._docx_path)

    def extract_images(self) -> Dict[str, str]:
        """
        Extrait les images du document .docx.

        Returns:
            Un dictionnaire {nom_image: chemin_relatif} ou {nom_image: donnees_base64}
        """
        logging.info("Extraction des images...")

        # S'assurer que le document est chargé
        if self._document is None:
            self.load_document()
            if self._document is None:
                logging.error("Impossible de charger le document")
                return {}

        rels = self._document.part.rels

        # Si on sauvegarde sur disque, créer le dossier images
        if self._save_images_to_disk:
            images_dir = os.path.join(self._output_dir, "images")
            if not os.path.exists(images_dir):
                os.makedirs(images_dir)
                logging.info(f"Dossier images créé: {images_dir}")

        for rel in rels.values():
            if "image" in rel.target_ref:
                try:
                    # Obtenir le nom de l'image depuis le chemin
                    image_name = os.path.basename(rel.target_ref)

                    # Vérifier si c'est une image externe
                    if rel.is_external:
                        logging.warning(f"Image externe ignorée: {rel.target_ref}")
                        continue

                    # Récupérer les données binaires
                    image_data = rel.target_part.blob

                    if self._save_images_to_disk:
                        # Chemin de sortie pour l'image
                        image_path = os.path.join(images_dir, image_name)

                        # Sauvegarder l'image
                        with open(image_path, "wb") as f:
                            f.write(image_data)

                        # Stocker le chemin relatif dans le dictionnaire
                        self._images[image_name] = f"images/{image_name}"
                        logging.debug(f"Image extraite et sauvegardée: {image_path}")
                    else:
                        # Encoder en base64
                        encoded_image = base64.b64encode(image_data).decode("utf-8")
                        # Stocker dans le dictionnaire
                        self._images[image_name] = encoded_image
                        logging.debug(
                            f"Image extraite et encodée en base64: {image_name}"
                        )

                except Exception as e:
                    logging.warning(
                        f"Impossible d'extraire l'image {rel.target_ref}: {str(e)}"
                    )

        # S'assurer que self._document n'est pas None avant d'accéder à part
        if self._document is not None:
            # Créer un dictionnaire des relations pour associer les rId aux chemins d'images
            for rel in self._document.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_name = os.path.basename(rel.target_ref)
                        if image_name in self._images:
                            if self._save_images_to_disk:
                                self._rels_dict[rel.rId] = self._images[image_name]
                            else:
                                # Pour les images en base64
                                self._rels_dict[rel.rId] = (
                                    f"data:image/png;base64,{self._images[image_name]}"
                                )
                    except Exception as e:
                        logging.warning(
                            f"Impossible de traiter la relation {rel.rId}: {str(e)}"
                        )

        return self._images

    def parse_paragraph(self, paragraph: DocxParagraph) -> DocumentElement:
        """
        Convertit un paragraphe en élément de document.

        Args:
            paragraph: L'objet Paragraph de python-docx

        Returns:
            Un élément de document (Paragraph, Heading, ListItem, etc.)
        """
        # Vérifier si le paragraphe contient une image
        for run in paragraph.runs:
            if hasattr(run, "_element"):
                # Recherche de l'élément inline (indiquant une image)
                inline_elem = None

                # Essayer différentes approches pour trouver les éléments d'image
                try:
                    # Méthode 1: Recherche directe avec le préfixe wp (si disponible)
                    if "wp" in run._element.nsmap:
                        inline_elem = run._element.find(
                            ".//wp:inline", namespaces=run._element.nsmap
                        )
                except Exception:
                    logging.debug(
                        "Impossible de trouver wp:inline avec les espaces de noms disponibles"
                    )

                # Si pas trouvé, essayer d'autres méthodes
                if inline_elem is None:
                    try:
                        # Méthode 2: Recherche pour n'importe quel élément inline sans préfixe spécifique
                        for child in run._element.iter(
                            "*"
                        ):  # Utiliser '*' pour tous les tags
                            if child.tag.endswith("inline"):
                                inline_elem = child
                                break
                    except Exception:
                        logging.debug(
                            "Impossible de rechercher un élément 'inline' générique"
                        )

                # Si on a trouvé une image
                if inline_elem is not None:
                    # Essayer de trouver l'attribut rId
                    rId = None
                    try:
                        # Rechercher l'élément blip et son attribut r:embed
                        for child in inline_elem.iter():
                            if child.tag.endswith("blip"):
                                # Chercher tous les attributs qui pourraient contenir l'ID de relation
                                for attrib_name, value in child.attrib.items():
                                    if attrib_name.endswith("embed"):
                                        rId = value
                                        break
                                if rId:
                                    break
                    except Exception as e:
                        logging.debug(
                            f"Impossible de récupérer l'ID de relation: {str(e)}"
                        )

                    # Créer l'élément image
                    image_element = Image(paragraph.text.strip() or "Image")

                    # Ajouter l'ID de relation s'il a été trouvé
                    if rId:
                        image_element.rId = rId
                        # Ajouter le chemin d'image s'il est disponible
                        if rId in self._rels_dict:
                            image_element.image_path = self._rels_dict[rId]

                    return image_element

        # Déterminer le type (paragraphe normal, titre, etc.)
        style_name = (
            paragraph.style.name if paragraph.style else ""
        )  # Utiliser une chaîne vide si None

        # Vérifier si c'est une instruction intégrée (commençant par ":::")
        text = paragraph.text.strip()
        if text.startswith(":::"):
            return Instruction(text[3:].strip())  # Retirer les ":::" du début

        # Vérifier si c'est un marqueur de composant pédagogique
        component_types = [
            "Vidéo",
            "Audio",
            "Accordéon",
            "Carrousel",
            "Onglets",
            "Défilement",
        ]
        if text.startswith("[") and any(
            text.startswith(f"[{comp}") for comp in component_types
        ):
            component_type = text[1:].split("]")[0] if "]" in text else "Unknown"
            return ComponentMarker(component_type)

        # Vérifier si c'est un marqueur de fin de composant
        if text.startswith("[Fin ") and "]" in text:
            component_type = text[5:].split("]")[0]
            return ComponentEnd(component_type)

        # Détecter si c'est un titre (Heading)
        if style_name and style_name.startswith("Heading"):
            level = (
                int(style_name.split(" ")[1]) if len(style_name.split(" ")) > 1 else 1
            )
            heading = Heading(level)

            # Extraire les runs (morceaux de texte avec style)
            for run in paragraph.runs:
                run_obj = TextRun(
                    text=run.text,
                    bold=bool(run.bold),  # Convertir en booléen strict
                    italic=bool(run.italic),  # Convertir en booléen strict
                    underline=bool(run.underline),  # Convertir en booléen strict
                )
                heading.add_run(run_obj)

            return heading

        # Vérifier si c'est un élément de liste
        if paragraph.paragraph_format.left_indent:
            list_item = ListItem()

            # Extraire les runs
            for run in paragraph.runs:
                run_obj = TextRun(
                    text=run.text,
                    bold=bool(run.bold),  # Convertir en booléen strict
                    italic=bool(run.italic),  # Convertir en booléen strict
                    underline=bool(run.underline),  # Convertir en booléen strict
                )
                list_item.add_run(run_obj)

            return list_item

        # C'est un paragraphe normal
        para = Paragraph()

        # Extraire les runs
        for run in paragraph.runs:
            run_obj = TextRun(
                text=run.text,
                bold=bool(run.bold),  # Convertir en booléen strict
                italic=bool(run.italic),  # Convertir en booléen strict
                underline=bool(run.underline),  # Convertir en booléen strict
            )
            para.add_run(run_obj)

        return para

    def parse_table(self, table: DocxTable) -> Table:
        """
        Convertit une table en élément Table.

        Args:
            table: L'objet Table de python-docx

        Returns:
            Un élément Table
        """
        table_element = Table()

        for row in table.rows:
            row_content = []
            for cell in row.cells:
                # Une cellule peut contenir plusieurs paragraphes
                cell_content = []
                for paragraph in cell.paragraphs:
                    para_element = self.parse_paragraph(paragraph)
                    # Ignorer les instructions dans les cellules
                    if para_element.type != "instruction":
                        cell_content.append(para_element)
                row_content.append(cell_content)
            table_element.add_row(row_content)

        return table_element

    def process_instructions(
        self, elements: List[DocumentElement]
    ) -> List[DocumentElement]:
        """
        Traite les instructions intégrées et les applique aux éléments suivants.

        Args:
            elements: Liste des éléments du document

        Returns:
            Liste des éléments modifiés selon les instructions
        """
        result = []
        current_block = None
        current_component = None
        skip_next = False

        # Classes et ID à appliquer au prochain élément
        pending_classes = []
        pending_id = None

        i = 0
        while i < len(elements):
            element = elements[i]

            # Si on doit ignorer cet élément (suite à une instruction "ignore")
            if skip_next:
                skip_next = False
                i += 1
                continue

            # Traitement des instructions spéciales :::
            if element.type == "instruction":
                # Vérifier que l'élément est bien une Instruction avant d'accéder à content
                if isinstance(element, Instruction):
                    instruction = element.content.strip()
                    logging.debug(f"Traitement de l'instruction: {instruction}")

                    # Instruction classe CSS
                    if instruction.startswith("class "):
                        class_names = instruction[6:].strip()
                        pending_classes.extend(class_names.split())

                    # Instruction ID
                    elif instruction.startswith("id "):
                        pending_id = instruction[3:].strip()

                    # Instruction ignore
                    elif instruction == "ignore":
                        skip_next = True

                    # Instruction bloc (pour wrapper des éléments)
                    elif " start" in instruction:
                        block_type = instruction.split()[0]
                        current_block = Block(block_type)

                    # Fin d'un bloc
                    elif " end" in instruction:
                        if current_block:
                            result.append(current_block)
                            current_block = None

                    # Instruction HTML brut
                    elif instruction.startswith("html "):
                        html_content = instruction[5:].strip()
                        result.append(RawHTML(html_content))

            # Traitement des marqueurs de composants pédagogiques
            elif element.type == "component_marker":
                # Vérifier que l'élément est bien un ComponentMarker avant d'accéder à component_type
                if isinstance(element, ComponentMarker):
                    component_type = element.component_type
                    logging.debug(f"Début de composant: {component_type}")

                    current_component = Component(component_type)

                    # Rechercher le contenu du composant jusqu'à [Fin] ou un autre composant
                    j = i + 1
                    while j < len(elements):
                        next_element = elements[j]

                        # Si on trouve un marqueur de fin pour ce composant
                        if (
                            next_element.type == "component_end"
                            and isinstance(next_element, ComponentEnd)
                            and next_element.component_type == component_type
                        ):
                            i = j  # On avance jusqu'à la fin du composant
                            break
                        # Si on trouve un autre marqueur de composant, on s'arrête aussi
                        elif next_element.type == "component_marker":
                            i = j - 1  # On s'arrête juste avant le nouveau composant
                            break
                        # Sinon, on ajoute l'élément au contenu du composant
                        else:
                            current_component.add_element(next_element)
                        j += 1

                    # Ajouter le composant au résultat
                    result.append(current_component)
                    current_component = None

            # Traitement des marqueurs de fin (sans marqueur de début correspondant)
            elif element.type == "component_end":
                # Vérifier que l'élément est bien un ComponentEnd avant d'accéder à component_type
                if isinstance(element, ComponentEnd):
                    logging.warning(
                        f"Marqueur de fin '{element.component_type}' sans marqueur de début correspondant"
                    )

            else:
                # Appliquer les classes et ID en attente
                if pending_classes:
                    element.html_class = " ".join(pending_classes)
                    pending_classes = []

                if pending_id:
                    element.html_id = pending_id
                    pending_id = None

                # Ajouter l'élément au bloc courant ou aux résultats
                if current_block:
                    current_block.add_element(element)
                elif current_component:
                    current_component.add_element(element)
                else:
                    result.append(element)

            i += 1

        # Ajouter le dernier bloc s'il existe encore
        if current_block:
            result.append(current_block)
            logging.warning("Un bloc n'a pas été fermé dans le document.")

        return result

    def convert(self) -> Dict[str, Any]:
        """
        Convertit le document DOCX en structure JSON.

        Returns:
            Un dictionnaire représentant le document
        """
        if self._document is None:
            self.load_document()
            if self._document is None:
                logging.error("Impossible de charger le document")
                return {
                    "meta": {
                        "title": os.path.basename(self._docx_path),
                        "error": "Document non chargé",
                    },
                    "content": [],
                    "images": {},
                }

        # Extraire les images si pas déjà fait
        if not self._images:
            self.extract_images()

        # Extraire le contenu en respectant l'ordre
        logging.info("Extraction du contenu...")
        elements = []

        # Vérifier que document est chargé et a un attribut element
        if hasattr(self._document, "element") and hasattr(
            self._document.element, "body"
        ):
            body = self._document.element.body

            # Parcourir les éléments du corps du document (paragraphes et tables)
            for child in body.iterchildren():
                if isinstance(child, CT_P):
                    # S'assurer que self._document n'est pas None pour le constructeur DocxParagraph
                    paragraph = DocxParagraph(child, self._document)
                    para_element = self.parse_paragraph(paragraph)
                    elements.append(para_element)
                elif isinstance(child, CT_Tbl):
                    # S'assurer que self._document n'est pas None pour le constructeur DocxTable
                    table = DocxTable(child, self._document)
                    table_element = self.parse_table(table)
                    elements.append(table_element)
        else:
            logging.error("Le document n'a pas de corps valide")
            return {
                "meta": {
                    "title": os.path.basename(self._docx_path),
                    "error": "Structure de document invalide",
                },
                "content": [],
                "images": self._images,
            }

        # Traiter les instructions intégrées
        logging.info("Traitement des instructions...")
        processed_elements = self.process_instructions(elements)

        # Créer la structure JSON complète
        file_name = os.path.basename(self._docx_path)
        result = {
            "meta": {"title": file_name},
            "content": [element.to_dict() for element in processed_elements],
            "images": self._images,
        }

        return result

    def generate_html(
        self, json_data: Dict[str, Any], css_path: Optional[str] = None
    ) -> str:
        """
        Génère un document HTML à partir de la structure JSON.

        Args:
            json_data: Dictionnaire représentant le document
            css_path: Chemin vers un fichier CSS personnalisé (optionnel)

        Returns:
            Une chaîne de caractères contenant le HTML
        """
        logging.info("Génération du HTML...")

        html_generator = HTMLGenerator(json_data)

        # Si un CSS personnalisé est fourni, le charger
        if css_path and os.path.exists(css_path):
            try:
                with open(css_path, "r", encoding="utf-8") as f:
                    custom_css = f.read()
                return html_generator.generate(custom_css=custom_css)
            except Exception as e:
                logging.warning(f"Impossible de charger le CSS personnalisé: {str(e)}")

        return html_generator.generate()

    def generate_markdown(self, json_data: Dict[str, Any]) -> str:
        """
        Génère un document Markdown à partir de la structure JSON.

        Args:
            json_data: Dictionnaire représentant le document

        Returns:
            Une chaîne de caractères contenant le Markdown
        """
        logging.info("Génération du Markdown...")

        md_generator = MarkdownGenerator(json_data)
        return md_generator.generate()


class HTMLGenerator:
    """
    Classe pour générer du HTML à partir de la structure JSON.
    """

    def __init__(self, json_data: Dict[str, Any]):
        """
        Initialise le générateur HTML.

        Args:
            json_data: Dictionnaire représentant le document
        """
        self._json_data = json_data
        self._images = json_data["images"]

    def generate(self, custom_css: Optional[str] = None) -> str:
        """
        Génère le document HTML complet.

        Args:
            custom_css: CSS personnalisé à utiliser au lieu du Bootstrap par défaut (optionnel)

        Returns:
            Une chaîne de caractères contenant le HTML
        """
        html = [
            "<!DOCTYPE html>",
            '<html lang="fr">',
            "<head>",
            f'<title>{self._json_data["meta"]["title"]}</title>',
            '<meta charset="UTF-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        ]

        # Utiliser le CSS personnalisé s'il est fourni, sinon Bootstrap
        if custom_css:
            html.append(f"<style>{custom_css}</style>")
        else:
            html.append(
                '<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-QWTKZyjpPEjISv5WaRU9OFeRpok6YctnYmDr5pNlyT2bRjXh0JMhjY6hW+ALEwIH" crossorigin="anonymous">'
            )

        html.append("</head>")
        html.append("<body>")
        html.append('<div class="container">')

        # Générer le HTML pour chaque élément
        for element in self._json_data["content"]:
            html.extend(self._generate_element_html(element))

        # Fin du document HTML
        html.append("</div>")  # Fermer le container

        # N'inclure le script Bootstrap que si on utilise Bootstrap (pas de CSS personnalisé)
        if not custom_css:
            html.append(
                '<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js" integrity="sha384-YvpcrYf0tY3lHB60NNkmXc5s9fDVZLESaAA55NDzOxhy9GkcIdslK1eN7N6jIeHz" crossorigin="anonymous"></script>'
            )

        html.append("</body>")
        html.append("</html>")

        return "\n".join(html)

    def _generate_element_html(self, element: Dict[str, Any]) -> List[str]:
        """
        Génère le HTML pour un élément spécifique.

        Args:
            element: Dictionnaire représentant l'élément

        Returns:
            Liste de chaînes de caractères HTML
        """
        element_html = []

        if element["type"] == "paragraph":
            attrs = []
            if "html_class" in element:
                attrs.append(f'class="{element["html_class"]}"')
            if "html_id" in element:
                attrs.append(f'id="{element["html_id"]}"')

            attrs_str = " ".join(attrs)
            attrs_str = f" {attrs_str}" if attrs_str else ""

            element_html.append(f"<p{attrs_str}>")
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"<strong>{text}</strong>"
                if run["italic"]:
                    text = f"<em>{text}</em>"
                if run["underline"]:
                    text = f"<u>{text}</u>"
                element_html.append(text)
            element_html.append("</p>")

        elif element["type"] == "heading":
            attrs = []
            if "html_class" in element:
                attrs.append(f'class="{element["html_class"]}"')
            if "html_id" in element:
                attrs.append(f'id="{element["html_id"]}"')

            attrs_str = " ".join(attrs)
            attrs_str = f" {attrs_str}" if attrs_str else ""

            level = element["level"]
            element_html.append(f"<h{level}{attrs_str}>")
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"<strong>{text}</strong>"
                if run["italic"]:
                    text = f"<em>{text}</em>"
                if run["underline"]:
                    text = f"<u>{text}</u>"
                element_html.append(text)
            element_html.append(f"</h{level}>")

        elif element["type"] == "list_item":
            # Note: ceci est simplifié et ne gère pas les listes imbriquées correctement
            element_html.append("<li>")
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"<strong>{text}</strong>"
                if run["italic"]:
                    text = f"<em>{text}</em>"
                if run["underline"]:
                    text = f"<u>{text}</u>"
                element_html.append(text)
            element_html.append("</li>")

        elif element["type"] == "table":
            element_html.append('<table class="table table-bordered">')
            for row in element["rows"]:
                element_html.append("<tr>")
                for cell in row:
                    element_html.append("<td>")
                    for para in cell:
                        element_html.extend(self._generate_element_html(para))
                    element_html.append("</td>")
                element_html.append("</tr>")
            element_html.append("</table>")

        elif element["type"] == "raw_html":
            element_html.append(element["content"])

        elif element["type"] == "block":
            block_type = element["block_type"]
            if block_type == "quote":
                element_html.append('<blockquote class="blockquote">')
                for content_elem in element["content"]:
                    element_html.extend(self._generate_element_html(content_elem))
                element_html.append("</blockquote>")
            elif block_type == "aside":
                element_html.append('<aside class="border p-3 my-3">')
                for content_elem in element["content"]:
                    element_html.extend(self._generate_element_html(content_elem))
                element_html.append("</aside>")
            # Ajoutez d'autres types de blocs au besoin

        elif element["type"] == "component":
            component_type = element["component_type"]

            if component_type == "Vidéo":
                # Template pour une vidéo (exemple)
                element_html.append('<div class="ratio ratio-16x9 my-4">')
                element_html.append(
                    '  <iframe src="https://www.youtube.com/embed/dQw4w9WgXcQ" title="Video" allowfullscreen></iframe>'
                )
                element_html.append("</div>")
                element_html.append('<div class="mt-2 mb-4">')
                for content_elem in element["content"]:
                    element_html.extend(self._generate_element_html(content_elem))
                element_html.append("</div>")

            elif component_type == "Audio":
                # Template pour audio
                element_html.append('<div class="my-4">')
                element_html.append('  <audio controls class="w-100">')
                element_html.append('    <source src="#" type="audio/mpeg">')
                element_html.append(
                    "    Votre navigateur ne supporte pas l'élément audio."
                )
                element_html.append("  </audio>")
                element_html.append("</div>")
                element_html.append('<div class="mt-2 mb-4">')
                for content_elem in element["content"]:
                    element_html.extend(self._generate_element_html(content_elem))
                element_html.append("</div>")

            elif component_type == "Accordéon":
                # Template pour accordéon
                accordion_id = f"accordion-{id(element)}"
                element_html.append(f'<div class="accordion my-4" id="{accordion_id}">')

                # Parcourir le contenu pour créer les éléments d'accordéon
                item_count = 0
                title = "Accordéon"
                content_html = []

                for content_elem in element["content"]:
                    if content_elem["type"] == "heading":
                        # Si on a déjà un titre et du contenu, créer un item d'accordéon
                        if item_count > 0 and content_html:
                            item_id = f"{accordion_id}-item-{item_count}"
                            element_html.append(f'  <div class="accordion-item">')
                            element_html.append(
                                f'    <h2 class="accordion-header" id="heading-{item_id}">'
                            )
                            element_html.append(
                                f'      <button class="accordion-button collapsed" type="button" data-bs-toggle="collapse" data-bs-target="#collapse-{item_id}" aria-expanded="false" aria-controls="collapse-{item_id}">'
                            )
                            element_html.append(f"        {title}")
                            element_html.append(f"      </button>")
                            element_html.append(f"    </h2>")
                            element_html.append(
                                f'    <div id="collapse-{item_id}" class="accordion-collapse collapse" aria-labelledby="heading-{item_id}" data-bs-parent="#{accordion_id}">'
                            )
                            element_html.append(f'      <div class="accordion-body">')
                            element_html.extend(content_html)
                            element_html.append(f"      </div>")
                            element_html.append(f"    </div>")
                            element_html.append(f"  </div>")

                            # Réinitialiser
                            content_html = []

                        # Nouveau titre d'accordéon
                        title = "".join([run["text"] for run in content_elem["runs"]])
                        item_count += 1
                    else:
                        # Ajouter au contenu en cours
                        content_html.extend(self._generate_element_html(content_elem))

                # Ajouter le dernier item s'il existe
                if item_count > 0 and content_html:
                    item_id = f"{accordion_id}-item-{item_count}"
                    element_html.append(f'  <div class="accordion-item">')
                    element_html.append(
                        f'    <h2 class="accordion-header" id="heading-{item_id}">'
                    )
                    element_html.append(
                        f'      <button class="accordion-button collapsed" type="button" data-bs-toggle="collapse" data-bs-target="#collapse-{item_id}" aria-expanded="false" aria-controls="collapse-{item_id}">'
                    )
                    element_html.append(f"        {title}")
                    element_html.append(f"      </button>")
                    element_html.append(f"    </h2>")
                    element_html.append(
                        f'    <div id="collapse-{item_id}" class="accordion-collapse collapse" aria-labelledby="heading-{item_id}" data-bs-parent="#{accordion_id}">'
                    )
                    element_html.append(f'      <div class="accordion-body">')
                    element_html.extend(content_html)
                    element_html.append(f"      </div>")
                    element_html.append(f"    </div>")
                    element_html.append(f"  </div>")

                element_html.append("</div>")

            elif component_type == "Carrousel":
                # Template pour carrousel
                carousel_id = f"carousel-{id(element)}"
                element_html.append(
                    f'<div id="{carousel_id}" class="carousel slide my-4" data-bs-ride="carousel">'
                )
                element_html.append('  <div class="carousel-inner">')

                # Ajouter les éléments du carrousel
                for i, content_elem in enumerate(element["content"]):
                    active = " active" if i == 0 else ""
                    element_html.append(f'    <div class="carousel-item{active}">')
                    element_html.extend(self._generate_element_html(content_elem))
                    element_html.append("    </div>")

                element_html.append("  </div>")
                element_html.append(
                    f'  <button class="carousel-control-prev" type="button" data-bs-target="#{carousel_id}" data-bs-slide="prev">'
                )
                element_html.append(
                    '    <span class="carousel-control-prev-icon" aria-hidden="true"></span>'
                )
                element_html.append(
                    '    <span class="visually-hidden">Précédent</span>'
                )
                element_html.append("  </button>")
                element_html.append(
                    f'  <button class="carousel-control-next" type="button" data-bs-target="#{carousel_id}" data-bs-slide="next">'
                )
                element_html.append(
                    '    <span class="carousel-control-next-icon" aria-hidden="true"></span>'
                )
                element_html.append('    <span class="visually-hidden">Suivant</span>')
                element_html.append("  </button>")
                element_html.append("</div>")

            elif component_type == "Onglets":
                # Template pour onglets
                tabs_id = f"tabs-{id(element)}"
                element_html.append('<div class="my-4">')
                element_html.append('  <ul class="nav nav-tabs" role="tablist">')

                # Créer les onglets
                tab_contents = []
                for i, content_elem in enumerate(element["content"]):
                    if content_elem["type"] == "heading":
                        tab_id = f"{tabs_id}-tab-{i}"
                        active = " active" if i == 0 else ""
                        selected = "true" if i == 0 else "false"

                        # Le titre de l'onglet
                        tab_title = "".join(
                            [run["text"] for run in content_elem["runs"]]
                        )
                        element_html.append(
                            f'    <li class="nav-item" role="presentation">'
                        )
                        element_html.append(
                            f'      <button class="nav-link{active}" id="{tab_id}-button" data-bs-toggle="tab" data-bs-target="#{tab_id}" type="button" role="tab" aria-controls="{tab_id}" aria-selected="{selected}">{tab_title}</button>'
                        )
                        element_html.append("    </li>")

                        # Préparer le contenu de l'onglet (à ajouter plus tard)
                        tab_contents.append((tab_id, active, []))
                    elif len(tab_contents) > 0:
                        # Ajouter au contenu du dernier onglet
                        tab_contents[-1][2].extend(
                            self._generate_element_html(content_elem)
                        )

                element_html.append("  </ul>")
                element_html.append('  <div class="tab-content">')

                # Ajouter le contenu des onglets
                for tab_id, active, content in tab_contents:
                    element_html.append(
                        f'    <div class="tab-pane fade show{active}" id="{tab_id}" role="tabpanel" aria-labelledby="{tab_id}-button">'
                    )
                    element_html.extend(content)
                    element_html.append("    </div>")

                element_html.append("  </div>")
                element_html.append("</div>")

            elif component_type == "Défilement":
                # Template pour scrollspy
                scrollspy_id = f"scrollspy-{id(element)}"
                element_html.append('<div class="row my-4">')
                element_html.append('  <div class="col-4">')
                element_html.append(
                    f'    <nav id="{scrollspy_id}" class="navbar navbar-light bg-light flex-column align-items-stretch p-3">'
                )
                element_html.append('      <nav class="nav nav-pills flex-column">')

                # Créer les liens de navigation
                nav_items = []
                for i, content_elem in enumerate(element["content"]):
                    if content_elem["type"] == "heading":
                        item_id = f"{scrollspy_id}-item-{i}"
                        item_title = "".join(
                            [run["text"] for run in content_elem["runs"]]
                        )
                        active = " active" if i == 0 else ""

                        element_html.append(
                            f'        <a class="nav-link{active}" href="#{item_id}">{item_title}</a>'
                        )
                        nav_items.append((item_id, content_elem, []))
                    elif len(nav_items) > 0:
                        # Ajouter au contenu de la dernière section
                        nav_items[-1][2].append(content_elem)

                element_html.append("      </nav>")
                element_html.append("    </nav>")
                element_html.append("  </div>")

                element_html.append('  <div class="col-8">')
                element_html.append(
                    f'    <div data-bs-spy="scroll" data-bs-target="#{scrollspy_id}" data-bs-offset="0" class="scrollspy-example-2" tabindex="0" style="height: 400px; overflow-y: scroll;">'
                )

                # Ajouter le contenu des sections
                for item_id, heading, content_elems in nav_items:
                    element_html.append(
                        f'      <h4 id="{item_id}">{" ".join([run["text"] for run in heading["runs"]])}</h4>'
                    )
                    for content_elem in content_elems:
                        element_html.extend(self._generate_element_html(content_elem))

                element_html.append("    </div>")
                element_html.append("  </div>")
                element_html.append("</div>")

            else:
                # Pour les autres types de composants, juste englober dans un div
                element_html.append(
                    f'<div class="component-{component_type.lower()} p-3 my-3 border">'
                )
                element_html.append(f"<h3>{component_type}</h3>")
                for content_elem in element["content"]:
                    element_html.extend(self._generate_element_html(content_elem))
                element_html.append("</div>")

        # Vérifier si l'élément a une référence d'image associée
        elif element["type"] == "image" and "image_path" in element:
            img_path = element["image_path"]
            img_alt = element.get("alt_text", "Image")
            element_html.append(
                f'<img src="{img_path}" alt="{img_alt}" class="img-fluid my-3" />'
            )

        return element_html


class MarkdownGenerator:
    """
    Classe pour générer du Markdown à partir de la structure JSON.
    """

    def __init__(self, json_data: Dict[str, Any]):
        """
        Initialise le générateur Markdown.

        Args:
            json_data: Dictionnaire représentant le document
        """
        self._json_data = json_data
        self._images = json_data["images"]
        self._list_stack = []  # Pour suivre les listes imbriquées
        self._component_level = 0  # Pour suivre l'imbrication des composants

    def generate(self) -> str:
        """
        Génère le document Markdown complet.

        Returns:
            Une chaîne de caractères contenant le Markdown
        """
        md = []

        # En-tête du document avec métadonnées YAML (optionnel)
        md.append("---")
        md.append(f"title: {self._json_data['meta']['title']}")
        md.append("author: Généré automatiquement")
        md.append("date: " + datetime.datetime.now().strftime("%Y-%m-%d"))
        md.append("---")
        md.append("")  # Ligne vide après les métadonnées

        # Générer le Markdown pour chaque élément
        for element in self._json_data["content"]:
            md.extend(self._generate_element_md(element))
            md.append("")  # Ligne vide entre les éléments pour la lisibilité

        return "\n".join(md)

    def _generate_element_md(self, element: Dict[str, Any]) -> List[str]:
        """
        Génère le Markdown pour un élément spécifique.

        Args:
            element: Dictionnaire représentant l'élément

        Returns:
            Liste de chaînes de caractères Markdown
        """
        element_md = []

        if element["type"] == "paragraph":
            # Construire le texte du paragraphe
            para_text = ""
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"**{text}**"
                if run["italic"]:
                    text = f"*{text}*"
                if run["underline"]:
                    # Markdown n'a pas de syntaxe pour souligné, utilisons HTML ou une autre approche
                    text = f"<u>{text}</u>"  # HTML dans Markdown
                para_text += text

            # Ajouter les attributs HTML en commentaire si présents
            attrs = []
            if "html_class" in element:
                attrs.append(f"class=\"{element['html_class']}\"")
            if "html_id" in element:
                attrs.append(f"id=\"{element['html_id']}\"")

            if attrs:
                element_md.append(f"<!-- {' '.join(attrs)} -->")

            element_md.append(para_text)

        elif element["type"] == "heading":
            level = element["level"]
            # Construire le texte du titre
            heading_text = ""
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"**{text}**"
                if run["italic"]:
                    text = f"*{text}*"
                if run["underline"]:
                    text = f"<u>{text}</u>"  # HTML dans Markdown
                heading_text += text

            # Ajouter les attributs HTML en commentaire si présents
            attrs = []
            if "html_class" in element:
                attrs.append(f"class=\"{element['html_class']}\"")
            if "html_id" in element:
                attrs.append(f"id=\"{element['html_id']}\"")

            if attrs:
                element_md.append(f"<!-- {' '.join(attrs)} -->")

            # Deux styles de titres en markdown: # et ======/-------
            if level == 1:
                element_md.append(heading_text)
                element_md.append("=" * len(heading_text))
            elif level == 2:
                element_md.append(heading_text)
                element_md.append("-" * len(heading_text))
            else:
                # Titres de niveau 3 et au-delà
                element_md.append("#" * level + " " + heading_text)

        elif element["type"] == "list_item":
            # Construire le texte de l'élément de liste
            item_text = ""
            for run in element["runs"]:
                text = run["text"]
                if run["bold"]:
                    text = f"**{text}**"
                if run["italic"]:
                    text = f"*{text}*"
                if run["underline"]:
                    text = f"<u>{text}</u>"  # HTML dans Markdown
                item_text += text

            # Par défaut, utiliser des listes à puces
            element_md.append(f"- {item_text}")

        elif element["type"] == "table":
            # Créer un tableau Markdown
            for row_idx, row in enumerate(element["rows"]):
                # Construire la ligne du tableau
                cells = []
                for cell in row:
                    # Joindre tout le contenu de la cellule
                    cell_text = ""
                    for para in cell:
                        para_md = self._generate_element_md(para)
                        cell_text += " ".join(para_md)
                    cells.append(cell_text.strip())

                element_md.append("| " + " | ".join(cells) + " |")

                # Ajouter la ligne de séparation après l'en-tête
                if row_idx == 0:
                    element_md.append("| " + " | ".join(["---"] * len(row)) + " |")

        elif element["type"] == "raw_html":
            # Garder le HTML brut tel quel
            element_md.append(element["content"])

        elif element["type"] == "block":
            block_type = element["block_type"]
            if block_type == "quote":
                # Citation en Markdown
                for content_elem in element["content"]:
                    content_md = self._generate_element_md(content_elem)
                    for line in content_md:
                        # Préfixer chaque ligne avec >
                        element_md.append(f"> {line}")
            elif block_type == "aside":
                # Pas d'équivalent direct pour aside en Markdown
                element_md.append(":::info")
                for content_elem in element["content"]:
                    element_md.extend(self._generate_element_md(content_elem))
                element_md.append(":::")

        elif element["type"] == "component":
            component_type = element["component_type"]

            # Noter le début du composant
            element_md.append(f"[{component_type}]")
            element_md.append("")  # Ligne vide pour lisibilité

            # Traiter le contenu du composant
            for content_elem in element["content"]:
                element_md.extend(self._generate_element_md(content_elem))

            # Fermer le composant
            element_md.append("")  # Ligne vide pour lisibilité
            element_md.append(f"[Fin {component_type}]")

        return element_md
