#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Module pour l'extraction et l'analyse des documents DOCX
--------------------------------------------------------
"""

import base64
import logging
import os
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from docx_json.models import (
    Block,
    Component,
    ComponentEnd,
    ComponentMarker,
    DocumentElement,
    DocumentList,
    Heading,
    Image,
    Instruction,
    ListItem,
    Paragraph,
    RawHTML,
    Table,
    TextRun,
)
from docx_json.models.instruction import InstructionType, extract_title_value


class DocxParser:
    """
    Classe pour l'extraction et l'analyse des éléments d'un document DOCX.
    """

    def __init__(
        self, docx_path: str, output_dir: str, save_images_to_disk: bool = True
    ):
        """
        Initialise le parser.

        Args:
            docx_path: Chemin vers le fichier .docx
            output_dir: Répertoire de sortie pour les fichiers générés
            save_images_to_disk: Si True, sauvegarde les images sur disque. Sinon, encode en base64.
        """
        self._docx_path = docx_path
        self._output_dir = output_dir
        self._save_images_to_disk = save_images_to_disk
        self._document = None
        self._images = {}
        self._rels_dict = {}
        # Initialiser les métadonnées du document
        self.metadata = {
            "title": os.path.basename(docx_path)
        }  # Par défaut, le nom du fichier
        self.NAMESPACES = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        }
        self._load_relationships()

    def _load_relationships(self) -> None:
        """
        Charge les relations entre les éléments du document DOCX.
        """
        logging.info("Chargement des relations...")
        self._document = Document(self._docx_path)
        rels = self._document.part.rels

        # Si on sauvegarde sur disque, créer le dossier images
        if self._save_images_to_disk:
            images_dir = os.path.join(self._output_dir, "images")
            if not os.path.exists(images_dir):
                os.makedirs(images_dir)
                logging.info(f"Dossier images créé: {images_dir}")

        for rel in rels.values():
            if "image" in rel.target_ref:
                try:
                    # Obtenir le nom de l'image depuis le chemin
                    image_name = os.path.basename(rel.target_ref)

                    # Vérifier si c'est une image externe
                    if rel.is_external:
                        logging.warning(f"Image externe ignorée: {rel.target_ref}")
                        continue

                    # Récupérer les données binaires
                    image_data = rel.target_part.blob

                    if self._save_images_to_disk:
                        # Chemin de sortie pour l'image
                        image_path = os.path.join(images_dir, image_name)

                        # Sauvegarder l'image
                        with open(image_path, "wb") as f:
                            f.write(image_data)

                        # Stocker le chemin relatif dans le dictionnaire
                        self._images[image_name] = f"images/{image_name}"
                        logging.debug(f"Image extraite et sauvegardée: {image_path}")
                    else:
                        # Encoder en base64
                        encoded_image = base64.b64encode(image_data).decode("utf-8")
                        # Stocker dans le dictionnaire
                        self._images[image_name] = encoded_image
                        logging.debug(
                            f"Image extraite et encodée en base64: {image_name}"
                        )

                except Exception as e:
                    logging.warning(
                        f"Impossible d'extraire l'image {rel.target_ref}: {str(e)}"
                    )

        # S'assurer que self._document n'est pas None avant d'accéder à part
        if self._document is not None:
            # Créer un dictionnaire des relations pour associer les rIds aux chemins d'images
            for rel in self._document.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_name = os.path.basename(rel.target_ref)
                        if image_name in self._images:
                            if self._save_images_to_disk:
                                self._rels_dict[rel.rId] = self._images[image_name]
                            else:
                                # Pour les images en base64
                                self._rels_dict[rel.rId] = (
                                    f"data:image/png;base64,{self._images[image_name]}"
                                )
                    except Exception as e:
                        logging.warning(
                            f"Impossible de traiter la relation {rel.rId}: {str(e)}"
                        )

    def _get_runs(self, paragraph: DocxParagraph) -> List[Dict[str, Any]]:
        """
        Extraire les runs (morceaux de texte avec style) d'un paragraphe.
        """
        runs = []
        for run in paragraph.runs:
            runs.append(
                {
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                }
            )
        return runs

    def _parse_image(self, paragraph: DocxParagraph) -> Dict[str, Any]:
        """
        Parse une image dans un paragraphe.
        """
        for run in paragraph.runs:
            if run._element is not None:
                # Vérifier si le run contient une image
                if (
                    run._element.find(".//pic:pic", namespaces=self.NAMESPACES)
                    is not None
                ):
                    # Extraire l'ID de l'image
                    blip = run._element.find(".//a:blip", namespaces=self.NAMESPACES)
                    if blip is not None and hasattr(blip, "attrib"):
                        rId = blip.attrib.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if rId:
                            return {
                                "type": "image",
                                "alt_text": paragraph.text.strip() or "Image",
                                "rId": rId,
                                "image_path": self._rels_dict.get(rId),
                            }

        return {"type": "image", "alt_text": paragraph.text.strip() or "Image"}

    def _group_consecutive_lists(
        self, elements: List[DocumentElement]
    ) -> List[DocumentElement]:
        """
        Regroupe les éléments de liste consécutifs en un seul DocumentList.
        """
        grouped_elements: List[DocumentElement] = []
        current_list: Optional[DocumentList] = None

        for element in elements:
            if isinstance(element, DocumentList):
                if current_list is None:
                    current_list = element
                else:
                    # Ajouter les éléments de la liste courante à la liste en cours
                    for item in element.items:
                        current_list.add_item(item)
            else:
                if current_list is not None:
                    grouped_elements.append(current_list)
                    current_list = None
                grouped_elements.append(element)

        # Ajouter la dernière liste si elle existe
        if current_list is not None:
            grouped_elements.append(current_list)

        return grouped_elements

    def parse_paragraph(self, paragraph: DocxParagraph) -> DocumentElement:
        """
        Convertit un paragraphe en élément de document.
        """
        # Vérifier si c'est un marqueur de composant
        text = paragraph.text.strip()
        if text.startswith("[") and text.endswith("]"):
            # Extraire le type de composant et les attributs HTML
            component_pattern: str = r"\[([\w\s]+)(?:\s+(.+))?\]"
            match: Optional[re.Match] = re.match(component_pattern, text)
            if match:
                component_type: str = match.group(1).strip()
                attributes_str: str = match.group(2) or ""

                # Vérifier si c'est un marqueur de début ou de fin de composant
                if component_type in [
                    "Vidéo",
                    "Audio",
                    "Accordéon",
                    "Carrousel",
                    "Onglets",
                    "Défilement",
                ]:
                    logging.debug(
                        f"Marqueur de début de composant détecté: {component_type}"
                    )

                    # Extraire les attributs si présents (format: attr='value' attr2="value2")
                    component_marker = ComponentMarker(component_type=component_type)

                    if attributes_str:
                        # Pattern pour extraire les attributs au format key='value' ou key="value"
                        attr_pattern = r"(\w+)=(['\"])(.*?)\2"
                        for attr_match in re.finditer(attr_pattern, attributes_str):
                            attr_name: str = attr_match.group(1)
                            attr_value: str = attr_match.group(3)
                            # Stocker l'attribut dans l'objet ComponentMarker
                            setattr(component_marker, attr_name, attr_value)
                            logging.debug(
                                f"  Attribut détecté: {attr_name}={attr_value}"
                            )

                    return component_marker
                elif component_type.startswith("Fin "):
                    # Extraire le type de composant sans le "Fin "
                    end_component_type: str = component_type[4:].strip()
                    logging.debug(
                        f"Marqueur de fin de composant détecté: {end_component_type}"
                    )
                    return ComponentEnd(component_type=end_component_type)

        # Vérifier si c'est une image
        if any(
            run._element.find(".//pic:pic", namespaces=self.NAMESPACES) is not None
            for run in paragraph.runs
        ):
            image_data: Dict[str, Any] = self._parse_image(paragraph)
            image = Image(alt_text=image_data["alt_text"])
            if "rId" in image_data:
                image.rId = image_data["rId"]
            if "image_path" in image_data:
                image.image_path = image_data["image_path"]
            return image

        # Extraire toutes les images intégrées dans les runs du paragraphe
        paragraph_images = []
        for run in paragraph.runs:
            for blip in run._element.findall(".//a:blip", namespaces=self.NAMESPACES):
                if (
                    blip is not None
                    and blip.get(
                        "{http://schemas.openxmlformats.org/drawingml/2006/main}embed"
                    )
                    is not None
                ):
                    image_id = blip.get(
                        "{http://schemas.openxmlformats.org/drawingml/2006/main}embed"
                    )
                    if image_id in self._rels_dict:
                        image_path = self._rels_dict[image_id]
                        paragraph_images.append(
                            {
                                "path": image_path,
                                "alt_text": "Image intégrée au paragraphe",
                            }
                        )

        # Vérifier si c'est une instruction
        if paragraph.text.strip().startswith("{{") and paragraph.text.strip().endswith(
            "}}"
        ):
            return Instruction(content=paragraph.text.strip()[2:-2])

        # Vérifier si c'est un titre
        if (
            paragraph.style is not None
            and paragraph.style.name is not None
            and paragraph.style.name.startswith("Heading")
        ):
            level = int(paragraph.style.name[-1])
            heading = Heading(level=level)
            for run in paragraph.runs:
                heading.add_run(
                    TextRun(
                        text=run.text,
                        bold=bool(run.bold),
                        italic=bool(run.italic),
                        underline=bool(run.underline),
                    )
                )

            # Ajouter les images associées au titre s'il y en a
            if paragraph_images:
                heading.add_attribute("images", paragraph_images)

            return heading

        # Si le paragraphe est vide, on évite de le créer
        if not paragraph.text.strip() and not paragraph_images:
            return Paragraph()

        # Vérifier si c'est un élément de liste
        if (
            paragraph.paragraph_format.left_indent is not None
            and paragraph.paragraph_format.left_indent > 0
        ):
            is_ordered = False
            if paragraph.style and paragraph.style.name == "List Number":
                is_ordered = True
            elif paragraph.text.strip() and re.match(r"^\d+\.", paragraph.text.strip()):
                is_ordered = True

            # Créer un DocumentList avec les attributs appropriés
            doc_list = DocumentList(ordered=is_ordered)

            # Créer un ListItem et y ajouter les runs
            list_item = ListItem()
            for run in paragraph.runs:
                list_item.add_run(
                    TextRun(
                        text=run.text,
                        bold=bool(run.bold),
                        italic=bool(run.italic),
                        underline=bool(run.underline),
                    )
                )

            # Ajouter l'élément à la liste
            doc_list.add_item(list_item)

            # Ajouter les images associées à l'élément de liste s'il y en a
            if paragraph_images:
                list_item.add_attribute("images", paragraph_images)

            return doc_list

        # Paragraphe normal
        para = Paragraph()
        for run in paragraph.runs:
            para.add_run(
                TextRun(
                    text=run.text,
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                    underline=bool(run.underline),
                )
            )

        # Ajouter les images associées au paragraphe s'il y en a
        if paragraph_images:
            para.add_attribute("images", paragraph_images)

        return para

    def parse_table(self, table: DocxTable) -> Table:
        """
        Parse une table Word et retourne un élément Table.

        Args:
            table: Table Word à parser

        Returns:
            Un élément Table
        """
        table_element = Table()

        for row in table.rows:
            row_content: List[List[DocumentElement]] = []
            for cell in row.cells:
                # Une cellule peut contenir plusieurs paragraphes
                cell_content: List[DocumentElement] = []
                for paragraph in cell.paragraphs:
                    para_element: DocumentElement = self.parse_paragraph(paragraph)
                    # Ignorer les instructions dans les cellules
                    if para_element is not None and not isinstance(
                        para_element, Instruction
                    ):
                        cell_content.append(para_element)
                row_content.append(cell_content)
            table_element.add_row(row_content)

        return table_element

    def _process_instruction(self, instruction_text: str) -> Optional[DocumentElement]:
        """
        Traite une instruction et retourne l'élément correspondant.

        Args:
            instruction_text: Texte de l'instruction (sans les ':::')

        Returns:
            L'élément correspondant ou None si l'instruction doit être ignorée
        """
        instruction_type = InstructionType.from_text(instruction_text)

        # Traiter l'instruction de titre
        if instruction_type == InstructionType.TITLE:
            title_value = extract_title_value(instruction_text)
            if title_value:
                # Stocker la valeur du titre dans les métadonnées
                self.metadata["title"] = title_value
                logging.info(f"Titre du document défini: {title_value}")
                return None  # Pas besoin de retourner d'élément pour cette instruction

        # Instruction inconnue ou non implémentée
        return Instruction(content=instruction_text)

    def parse(self) -> List[DocumentElement]:
        """
        Parse un document Word et retourne une liste d'éléments de document.
        """
        if not self._document:
            raise ValueError("No document loaded")

        logging.info(f"Chargement du document: {self._docx_path}")
        logging.info("Extraction du contenu...")

        elements: List[DocumentElement] = []
        for paragraph in self._document.paragraphs:
            # Vérifier si c'est une instruction (texte commençant par ":::")
            text = paragraph.text.strip()
            if text.startswith(":::"):
                instruction_text = text[3:].strip()
                instruction_element = self._process_instruction(instruction_text)
                if instruction_element:
                    elements.append(instruction_element)
                continue

            # Sinon, traiter normalement
            para_element: DocumentElement = self.parse_paragraph(paragraph)
            if para_element is not None and not isinstance(para_element, Instruction):
                elements.append(para_element)

        # Regrouper les éléments de liste consécutifs
        elements = self._group_consecutive_lists(elements)

        return elements

    def get_images(self) -> Dict[str, str]:
        """
        Renvoie le dictionnaire des images extraites.

        Returns:
            Dictionnaire des images extraites
        """
        return self._images
