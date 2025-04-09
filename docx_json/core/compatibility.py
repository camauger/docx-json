#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Couche de compatibilité pour l'API originale
-------------------------------------------

Ce module fournit une compatibilité avec l'ancienne API tout en utilisant
la nouvelle implémentation orientée objet.
"""

import json
import logging
import os
import sys
from typing import Any, Dict, List, Optional, cast

from docx.table import Table
from docx.text.paragraph import Paragraph

from docx_json.core.converter import DocxConverter
from docx_json.core.html_renderer import HTMLGenerator
from docx_json.exceptions import DocxValidationError
from docx_json.models.elements import (
    Block,
    Component,
    ComponentEnd,
    ComponentMarker,
    DocumentElement,
    Heading,
    Image,
    Instruction,
    ListItem,
)
from docx_json.models.elements import Paragraph as ParagraphElement
from docx_json.models.elements import RawHTML
from docx_json.models.elements import Table as TableElement
from docx_json.models.elements import TextRun
from docx_json.utils.image_utils import extract_images as extract_images_util


def get_paragraph_json(paragraph: Paragraph) -> Dict[str, Any]:
    """
    Convertit un paragraphe en structure JSON

    Args:
        paragraph: L'objet Paragraph de python-docx

    Returns:
        Un dictionnaire représentant le paragraphe
    """
    # Créer un convertisseur temporaire
    converter = DocxConverter("", ".")
    # Analyser le paragraphe
    element = converter.parse_paragraph(paragraph)
    # Convertir en dictionnaire
    return element.to_dict()


def get_table_json(table: Table) -> Dict[str, Any]:
    """
    Convertit une table en structure JSON

    Args:
        table: L'objet Table de python-docx

    Returns:
        Un dictionnaire représentant la table
    """
    # Créer un convertisseur temporaire
    converter = DocxConverter("", ".")
    # Analyser la table
    table_element = converter.parse_table(table)
    # Convertir en dictionnaire
    return table_element.to_dict()


def create_element_from_dict(data: Dict[str, Any]) -> DocumentElement:
    """
    Crée un objet DocumentElement à partir d'un dictionnaire

    Args:
        data: Dictionnaire contenant les données de l'élément

    Returns:
        Un objet DocumentElement correspondant
    """
    element_type = data.get("type", "")

    if element_type == "paragraph":
        elem = ParagraphElement()
        # Ajouter les runs
        for run_data in data.get("runs", []):
            elem.add_run(
                TextRun(
                    text=run_data.get("text", ""),
                    bold=run_data.get("bold", False),
                    italic=run_data.get("italic", False),
                    underline=run_data.get("underline", False),
                )
            )

    elif element_type == "heading":
        elem = Heading(level=data.get("level", 1))
        # Ajouter les runs
        for run_data in data.get("runs", []):
            elem.add_run(
                TextRun(
                    text=run_data.get("text", ""),
                    bold=run_data.get("bold", False),
                    italic=run_data.get("italic", False),
                    underline=run_data.get("underline", False),
                )
            )

    elif element_type == "list_item":
        elem = ListItem()
        # Ajouter les runs
        for run_data in data.get("runs", []):
            elem.add_run(
                TextRun(
                    text=run_data.get("text", ""),
                    bold=run_data.get("bold", False),
                    italic=run_data.get("italic", False),
                    underline=run_data.get("underline", False),
                )
            )

    elif element_type == "image":
        elem = Image(alt_text=data.get("alt_text", "Image"))
        if "rId" in data:
            elem.rId = data["rId"]
        if "image_path" in data:
            elem.image_path = data["image_path"]

    elif element_type == "table":
        elem = TableElement()
        # Convertir les lignes (structure complexe)
        # Pour simplifier, on crée seulement un squelette de table ici
        # En pratique, il faudrait traiter récursivement les cellules

    elif element_type == "block":
        elem = Block(block_type=data.get("block_type", ""))
        # Ajouter les éléments du contenu
        for content_item in data.get("content", []):
            elem.add_element(create_element_from_dict(content_item))

    elif element_type == "component":
        elem = Component(component_type=data.get("component_type", ""))
        # Ajouter les éléments du contenu
        for content_item in data.get("content", []):
            elem.add_element(create_element_from_dict(content_item))

    elif element_type == "component_marker":
        elem = ComponentMarker(component_type=data.get("component_type", ""))

    elif element_type == "component_end":
        elem = ComponentEnd(component_type=data.get("component_type", ""))

    elif element_type == "instruction":
        elem = Instruction(content=data.get("content", ""))

    elif element_type == "raw_html":
        elem = RawHTML(content=data.get("content", ""))

    else:
        # Type inconnu, créer un élément de base
        elem = DocumentElement(element_type=element_type)

    # Appliquer les classes et ID en utilisant les setters qui gèrent les non-strings
    if "html_class" in data:
        # DocumentElement.html_class.setter accepte uniquement str, pas None
        if data["html_class"] is not None:
            # On utilise cast pour dire à mypy que ce n'est pas None
            elem.html_class = cast(str, data["html_class"])
    if "html_id" in data:
        # DocumentElement.html_id.setter accepte uniquement str, pas None
        if data["html_id"] is not None:
            # On utilise cast pour dire à mypy que ce n'est pas None
            elem.html_id = cast(str, data["html_id"])

    return elem


def process_instructions(elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Traite les instructions intégrées et les applique aux éléments suivants

    Args:
        elements: Liste des éléments JSON du document

    Returns:
        Liste des éléments modifiés selon les instructions
    """
    # Créer un convertisseur temporaire
    converter = DocxConverter("", ".")

    # Convertir les dictionnaires en objets DocumentElement
    element_objects = [create_element_from_dict(elem) for elem in elements]

    # Traiter les instructions
    processed_objects = converter.process_instructions(element_objects)

    # Convertir les objets en dictionnaires
    return [elem.to_dict() for elem in processed_objects]


def extract_images(
    document: Any, output_dir: str, save_to_disk: bool = True
) -> Dict[str, str]:
    """
    Extrait les images du document .docx et les sauvegarde dans un dossier ou les encode en base64

    Args:
        document: L'objet Document de python-docx
        output_dir: Répertoire de sortie pour les fichiers générés
        save_to_disk: Si True, sauvegarde les images sur disque. Sinon, encode en base64.

    Returns:
        Un dictionnaire {nom_image: chemin_relatif} ou {nom_image: donnees_base64}
    """
    # Déléguer à la fonction du module utils
    return extract_images_util(document, output_dir, save_to_disk)


def get_document_json(
    docx_path: str, output_dir: Optional[str] = None, save_images_to_disk: bool = True
) -> Dict[str, Any]:
    """
    Convertit un document .docx complet en structure JSON

    Args:
        docx_path: Chemin vers le fichier .docx
        output_dir: Répertoire de sortie pour les fichiers générés
        save_images_to_disk: Si True, sauvegarde les images sur disque. Sinon, encode en base64.

    Returns:
        Un dictionnaire représentant le document
    """
    # Si output_dir n'est pas spécifié, utiliser le répertoire du fichier docx
    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(docx_path))
        if not output_dir:
            output_dir = "."

    # Utiliser la nouvelle classe DocxConverter
    converter = DocxConverter(docx_path, output_dir, save_images_to_disk)
    return converter.convert()


def generate_html(json_data: Dict[str, Any], css_path: Optional[str] = None) -> str:
    """
    Génère une représentation HTML à partir de la structure JSON du document

    Args:
        json_data: Dictionnaire représentant le document
        css_path: Chemin vers un fichier CSS personnalisé (optionnel)

    Returns:
        Une chaîne de caractères contenant le HTML
    """
    # Charger le CSS personnalisé s'il est spécifié
    custom_css = None
    if css_path:
        try:
            with open(css_path, "r", encoding="utf-8") as f:
                custom_css = f.read()
        except Exception as e:
            logging.warning(f"Impossible de charger le CSS personnalisé: {str(e)}")

    # Utiliser le générateur HTML
    generator = HTMLGenerator(json_data)
    return generator.generate(custom_css)


def generate_markdown(json_data: Dict[str, Any]) -> str:
    """
    Génère un document Markdown à partir de la structure JSON

    Args:
        json_data: Dictionnaire représentant le document

    Returns:
        Une chaîne de caractères contenant le Markdown
    """
    # Utiliser la classe DocxConverter
    converter = DocxConverter("", ".")
    return converter.generate_markdown(json_data)


def validate_docx(docx_path: str) -> bool:
    """
    Validates if a file is a valid DOCX document.

    Args:
        docx_path: Path to the DOCX file to validate

    Returns:
        bool: True if the file is a valid DOCX document, False otherwise

    Raises:
        DocxValidationError: When the file doesn't exist, has wrong extension, or can't be opened
    """
    if not os.path.exists(docx_path):
        raise DocxValidationError(f"File not found: {docx_path}")

    if not docx_path.lower().endswith(".docx"):
        raise DocxValidationError(f"Invalid file extension: {docx_path}")

    try:
        from docx import Document

        Document(docx_path)
        return True
    except Exception as e:
        raise DocxValidationError(f"Invalid DOCX file: {e}") from e
