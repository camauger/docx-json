#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Script pour tester les composants Audio et Vidéo dans le convertisseur docx-json.
Génère un fichier docx avec les balises appropriées pour les composants.
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Créer un nouveau document
doc = Document()

# Configurer le style des titres
doc.add_heading("Test des composants Audio et Vidéo", level=1)

# Ajouter un composant Audio
doc.add_heading("Composant Audio", level=2)
doc.add_paragraph("Voici un exemple de composant audio:")

# Ajouter les marqueurs du composant Audio
audio_start = doc.add_paragraph("[Audio]")
doc.add_paragraph("Découvrez les explications de Marie en écoutant l'extrait suivant.")
doc.add_paragraph("Marie, représentante du SASIE")
audio_end = doc.add_paragraph("[Fin Audio]")

# Vérifier le texte exact des marqueurs
print(f"Marqueur de début Audio: '{audio_start.text}'")
print(f"Marqueur de fin Audio: '{audio_end.text}'")

# Ajouter un composant Vidéo simple
doc.add_heading("Composant Vidéo Simple", level=2)
doc.add_paragraph("Voici un exemple de composant vidéo simple:")

# Ajouter les marqueurs du composant Vidéo simple
video_start = doc.add_paragraph("[Vidéo]")
doc.add_paragraph("Une vidéo simple sans attributs.")
video_end = doc.add_paragraph("[Fin Vidéo]")

# Vérifier le texte exact des marqueurs
print(f"Marqueur de début Vidéo simple: '{video_start.text}'")
print(f"Marqueur de fin Vidéo simple: '{video_end.text}'")

# Ajouter un composant Vidéo avec ID
doc.add_heading("Composant Vidéo avec ID", level=2)
doc.add_paragraph("Voici un exemple de composant vidéo avec ID:")

# Ajouter les marqueurs du composant Vidéo avec ID
video_id_start = doc.add_paragraph('[Vidéo video_id="1069341210"]')
doc.add_paragraph("Regardez cette vidéo explicative.")
video_id_end = doc.add_paragraph("[Fin Vidéo]")

# Vérifier le texte exact des marqueurs
print(f"Marqueur de début Vidéo avec ID: '{video_id_start.text}'")
print(f"Marqueur de fin Vidéo avec ID: '{video_id_end.text}'")

# Ajouter un composant Consignes
doc.add_heading("Composant Consignes", level=2)
doc.add_paragraph("Voici un exemple de composant consignes:")

# Ajouter les marqueurs du composant Consignes
consignes_start = doc.add_paragraph("[Consignes]")
doc.add_paragraph("Cliquez sur le bouton pour continuer.")
doc.add_paragraph("N'oubliez pas de compléter tous les champs obligatoires.")
consignes_end = doc.add_paragraph("[Fin Consignes]")

# Vérifier le texte exact des marqueurs
print(f"Marqueur de début Consignes: '{consignes_start.text}'")
print(f"Marqueur de fin Consignes: '{consignes_end.text}'")

# Enregistrer le document
fichier_sortie = "audio_video_test.docx"
doc.save(fichier_sortie)
print(f"\nFichier '{fichier_sortie}' créé avec succès!")
print("Exécutez la commande suivante pour tester la conversion:")
print(f"python -m docx_json {fichier_sortie} --html --json --verbose")
