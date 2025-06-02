from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Créer un nouveau document
doc = Document()

# Ajouter le marqueur de début de carrousel
doc.add_paragraph("[Carrousel]")

# Ajouter trois diapositives contenant des balises d'images
for i in range(3):
    # Ajouter du texte qui sera reconnu comme une balise d'image
    paragraph = doc.add_paragraph(f"![Image {i+1}](images/placeholder.jpg)")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Ajouter le marqueur de fin de carrousel
doc.add_paragraph("[Fin Carrousel]")

# Sauvegarder le document
doc.save("test_templates/test_carousel_img.docx")

print("Fichier créé: test_templates/test_carousel_img.docx")
