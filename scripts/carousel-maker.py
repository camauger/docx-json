from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Créer un nouveau document
doc = Document()

# Ajouter le marqueur de début de carrousel
doc.add_paragraph("[Carrousel]")

# Ajouter trois diapositives avec du texte
for i in range(3):
    # Ajouter un titre pour chaque diapositive
    heading = doc.add_paragraph(f"Diapositive {i+1}")
    heading.style = "Heading 3"

    # Ajouter du contenu
    paragraph = doc.add_paragraph(
        f"Contenu de la diapositive {i+1}. Cette diapositive fera partie du carrousel."
    )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Ajouter le marqueur de fin de carrousel
doc.add_paragraph("[Fin Carrousel]")

# Sauvegarder le document
doc.save("test_templates/test_carousel.docx")

print("Fichier créé: test_templates/test_carousel.docx")
