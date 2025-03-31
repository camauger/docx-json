#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Tests unitaires pour la couche de compatibilité docx_json.core.compatibility

Ce fichier teste les fonctions de compatibilité qui fournissent une interface
vers l'implémentation orientée objet sous-jacente.
"""

import json
import os
import tempfile
import unittest

from docx import Document

from docx_json.core.compatibility import (
    generate_html,
    get_paragraph_json,
    get_table_json,
    process_instructions,
)


class TestConvertFunctions(unittest.TestCase):
    """Tests unitaires pour les fonctions de conversion principales."""

    def setUp(self):
        """Initialise l'environnement de test."""
        # Créer un document temporaire pour les tests
        self.test_doc = Document()
        self.test_doc.add_heading("Titre de test", level=1)
        self.test_doc.add_paragraph("Ceci est un paragraphe de test.")

        # Créer un tableau simple
        table = self.test_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        # Enregistrer le document
        self.temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.test_doc.save(self.temp_file.name)
        self.temp_file.close()

    def tearDown(self):
        """Nettoie l'environnement après les tests."""
        if os.path.exists(self.temp_file.name):
            os.unlink(self.temp_file.name)

    def test_paragraph_json(self):
        """Teste la conversion d'un paragraphe en JSON."""
        para = self.test_doc.paragraphs[1]  # Le deuxième paragraphe (après le titre)
        result = get_paragraph_json(para)

        self.assertEqual(result["type"], "paragraph")
        self.assertTrue(isinstance(result["runs"], list))
        self.assertEqual(result["runs"][0]["text"], "Ceci est un paragraphe de test.")

    def test_heading_json(self):
        """Teste la conversion d'un titre en JSON."""
        heading = self.test_doc.paragraphs[0]  # Le premier paragraphe (titre)
        result = get_paragraph_json(heading)

        self.assertEqual(result["type"], "heading")
        self.assertEqual(result["level"], 1)
        self.assertTrue("runs" in result)

    def test_table_json(self):
        """Teste la conversion d'un tableau en JSON."""
        table = self.test_doc.tables[0]
        result = get_table_json(table)

        self.assertEqual(result["type"], "table")
        self.assertEqual(len(result["rows"]), 2)
        self.assertEqual(len(result["rows"][0]), 2)

        # Vérifier le contenu des cellules
        cell_text = result["rows"][0][0][0]["runs"][0]["text"]
        self.assertEqual(cell_text, "A1")

    def test_process_instructions(self):
        """Teste le traitement des instructions intégrées."""
        elements = [
            {"type": "instruction", "content": "class hero dark"},
            {
                "type": "heading",
                "level": 1,
                "runs": [
                    {"text": "Titre", "bold": True, "italic": False, "underline": False}
                ],
            },
            {"type": "instruction", "content": "quote start"},
            {
                "type": "paragraph",
                "runs": [
                    {
                        "text": "Citation",
                        "bold": False,
                        "italic": False,
                        "underline": False,
                    }
                ],
            },
            {"type": "instruction", "content": "quote end"},
        ]

        result = process_instructions(elements)

        # Vérifier que l'élément titre a bien reçu les classes
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0]["html_class"], "hero dark")

        # Vérifier que le bloc quote contient bien le paragraphe
        self.assertEqual(result[1]["type"], "block")
        self.assertEqual(result[1]["block_type"], "quote")
        self.assertEqual(len(result[1]["content"]), 1)

    def test_generate_html(self):
        """Teste la génération de HTML à partir d'une structure JSON."""
        json_data = {
            "meta": {"title": "test.docx"},
            "content": [
                {
                    "type": "heading",
                    "level": 1,
                    "runs": [
                        {
                            "text": "Titre",
                            "bold": True,
                            "italic": False,
                            "underline": False,
                        }
                    ],
                },
                {
                    "type": "paragraph",
                    "runs": [
                        {
                            "text": "Paragraphe",
                            "bold": False,
                            "italic": False,
                            "underline": False,
                        }
                    ],
                },
            ],
            "images": {},
        }

        html = generate_html(json_data)

        self.assertTrue("<!DOCTYPE html>" in html)
        self.assertTrue("<h1>" in html)
        self.assertTrue("<strong>Titre</strong>" in html)
        self.assertTrue("<p>" in html)


if __name__ == "__main__":
    unittest.main()
