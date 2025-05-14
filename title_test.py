#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Crée un document Word avec un titre personnalisé et un contenu simple
pour tester l'instruction :::title.
"""

import os

from docx import Document


def create_title_test_document():
    # Créer un nouveau document Word
    doc = Document()

    # Ajouter l'instruction de titre
    p = doc.add_paragraph(":::title Mon Document avec Titre Personnalisé")

    # Ajouter un titre et du contenu normal
    h = doc.add_heading("Titre de section", level=1)
    p = doc.add_paragraph(
        "Ceci est un paragraphe de test pour vérifier que l'instruction :::title fonctionne correctement."
    )

    # Ajouter une autre instruction de classe
    p = doc.add_paragraph(":::class hero")
    p = doc.add_paragraph("Ceci est un texte qui devrait avoir la classe 'hero'")

    # Sauvegarder le document
    output_path = "titre_test.docx"
    doc.save(output_path)
    print(f"Document créé avec succès : {os.path.abspath(output_path)}")


if __name__ == "__main__":
    create_title_test_document()
