#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Script pour tester les composants Consignes dans le convertisseur docx-json.
Génère un fichier docx avec les balises appropriées pour ce composant.
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Créer un nouveau document
doc = Document()

# Configurer le style des titres
doc.add_heading("Test du composant Consignes", level=1)

# Ajouter un composant Consignes
doc.add_heading("Exemple de composant Consignes", level=2)
doc.add_paragraph("Voici un exemple de composant consignes:")

# Ajouter les marqueurs du composant Consignes
# Créer un paragraphe pour le marqueur de début
p_start = doc.add_paragraph()
p_start.add_run("[Consignes]")

# Ajouter le contenu du composant
doc.add_paragraph("Cliquez sur le bouton pour continuer.")
doc.add_paragraph("N'oubliez pas de compléter tous les champs obligatoires.")

# Créer un paragraphe pour le marqueur de fin
p_end = doc.add_paragraph()
p_end.add_run("[Fin Consignes]")

# Enregistrer le document
fichier_sortie = "consignes_simple_test.docx"
doc.save(fichier_sortie)
print(f"Fichier '{fichier_sortie}' créé avec succès!")
print("Exécutez la commande suivante pour tester la conversion:")
print(f"python -m docx_json {fichier_sortie} --html --json --verbose")
